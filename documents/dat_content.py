#!/usr/bin/env python3
"""Contenu du DAT (sections 1 a 23). Importe et assemble par generate_dat.py.

Chaque fonction add_section_XX(doc) ajoute une section complete au document
docx en cours de construction. Les faits presentes proviennent de
l'exploration verifiee du code source (voir recherche menee en amont) :
backend Spring Boot, frontend React, service DLP FastAPI, migrations
Flyway, Keycloak, LiteLLM, Docker Compose, tests.
"""

from pathlib import Path
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from generate_dat import (
    add_table, add_figure, add_page_break, DIAGRAMS_DIR, NAVY, ACCENT, GREY, WARN, GREEN,
)

D = DIAGRAMS_DIR


def p(doc, text, bold=False, italic=False, size=None, color=None, style=None, space_after=None):
    para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    return para


def bullets(doc, items, style="List Bullet"):
    for it in items:
        doc.add_paragraph(it, style=style)


def note(doc, text, label="Remarque"):
    para = doc.add_paragraph()
    r1 = para.add_run(f"{label} — ")
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = WARN
    r2 = para.add_run(text)
    r2.italic = True
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = GREY


# ---------------------------------------------------------------------------
# 1. Introduction
# ---------------------------------------------------------------------------

def add_section_01(doc):
    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("1.1 Contexte du projet", level=2)
    p(doc, "Secure LLM Gateway (nom de code interne du realm/theme : « Synapse ») est une "
           "passerelle applicative qui permet a des utilisateurs authentifies d'echanger avec "
           "plusieurs fournisseurs de modeles de langage (LLM) externes (OpenAI, Groq, Google "
           "Gemini, Mistral) sans jamais leur envoyer directement des donnees sensibles non "
           "controlees. Toute requete transite par un backend applicatif qui authentifie "
           "l'utilisateur, applique des regles d'autorisation, et fait analyser le contenu par un "
           "moteur de prevention de perte de donnees (DLP) avant tout appel au modele.")

    doc.add_heading("1.2 Objectif du projet", level=2)
    p(doc, "Fournir un point d'acces unique et controle aux LLM externes, avec authentification "
           "centralisee (Keycloak), autorisation par role, filtrage/masquage des donnees "
           "sensibles (DLP), tracabilite des actions administratives (audit) et une interface "
           "d'administration pour piloter les modeles, les utilisateurs, les roles et les regles "
           "de securite sans intervention directe en base de donnees.")

    doc.add_heading("1.3 Objectif du DAT", level=2)
    p(doc, "Ce document decrit l'architecture reelle du systeme telle qu'elle existe dans le code "
           "source, la configuration et l'infrastructure du projet au moment de sa redaction. Il "
           "ne documente aucune fonctionnalite hypothetique ou planifiee : chaque affirmation est "
           "verifiable dans le depot Git (backend, frontend, service DLP, migrations, Docker, "
           "Keycloak, LiteLLM, tests).")

    doc.add_heading("1.4 Perimetre", level=2)
    bullets(doc, [
        "Backend Spring Boot (API REST + SSE, package com.example.backend).",
        "Frontend React/Vite (application monopage, sans routeur URL).",
        "Service DLP Python/FastAPI (detection et masquage de donnees sensibles).",
        "Base de donnees PostgreSQL et migrations Flyway (V1 a V22).",
        "Authentification/autorisation Keycloak (realm synapse).",
        "Proxy LiteLLM vers les fournisseurs LLM externes.",
        "Infrastructure Docker Compose et integration continue GitHub Actions.",
        "Suites de tests automatisees des trois services applicatifs.",
    ])

    doc.add_heading("1.5 Public cible", level=2)
    p(doc, "Equipe de developpement et de maintenance du projet, evaluateurs academiques, "
           "responsables securite/architecture souhaitant comprendre le fonctionnement reel du "
           "systeme avant une revue ou une evolution.")

    doc.add_heading("1.6 References", level=2)
    bullets(doc, [
        "README.md (racine du depot) — instructions d'installation et d'exploitation.",
        "AUDIT_MERGE_INTEGRATION.md — audit de l'integration authentification/securite/DLP dans main.",
        "dlp/README.md — documentation du service DLP.",
        "litellm/README.md — documentation d'usage du proxy LiteLLM (a noter : alias Groq desynchronise, voir section 22).",
        "database/README.md — documentation de la base (partiellement obsolete, voir section 22).",
        "Code source du depot (source de verite prioritaire en cas de contradiction avec toute documentation).",
    ])

    doc.add_heading("1.7 Glossaire et acronymes", level=2)
    add_table(doc, ["Terme", "Definition"], [
        ["DAT", "Dossier d'Architecture Technique"],
        ["DLP", "Data Loss Prevention — prevention de la perte/fuite de donnees sensibles"],
        ["LLM", "Large Language Model — modele de langage de grande taille"],
        ["JWT", "JSON Web Token — jeton d'authentification signe utilise par Keycloak"],
        ["OIDC", "OpenID Connect — protocole d'authentification base sur OAuth2"],
        ["PKCE", "Proof Key for Code Exchange — extension de securite du flux OAuth2 Authorization Code"],
        ["SSE", "Server-Sent Events — flux HTTP unidirectionnel serveur vers client, utilise pour le streaming des reponses LLM"],
        ["NER", "Named Entity Recognition — reconnaissance d'entites nommees (utilisee par Presidio)"],
        ["CIN", "Carte d'Identite Nationale (marocaine) — type de donnee sensible detecte par le DLP"],
        ["RIB / IBAN", "Identifiants bancaires detectes par le DLP"],
        ["Flyway", "Outil de migration de schema de base de donnees versionnee"],
        ["Realm (Keycloak)", "Espace isole de gestion des utilisateurs/roles/clients dans Keycloak (ici : synapse)"],
    ], col_widths=[3.5, 12.5])


# ---------------------------------------------------------------------------
# 2. Presentation generale du systeme
# ---------------------------------------------------------------------------

def add_section_02(doc):
    doc.add_heading("2. Presentation generale du systeme", level=1)

    doc.add_heading("2.1 Problematique", level=2)
    p(doc, "L'usage direct de services LLM externes par des utilisateurs internes pose un risque "
           "de fuite de donnees sensibles (identifiants, secrets techniques, donnees personnelles) "
           "vers des tiers, ainsi qu'un manque de tracabilite et de controle d'acces par profil "
           "d'utilisateur.")

    doc.add_heading("2.2 Objectif de la Gateway", level=2)
    p(doc, "Centraliser tous les echanges avec les LLM derriere un backend qui : authentifie "
           "l'utilisateur (Keycloak/JWT), verifie ses autorisations (role et restrictions "
           "individuelles/par role), fait analyser chaque message et chaque piece jointe par le "
           "service DLP avant tout envoi au modele, journalise les incidents et les actions "
           "d'administration, et propose une interface d'administration pour piloter le catalogue "
           "de modeles, les utilisateurs Keycloak, les roles, les restrictions et les regles DLP.")

    doc.add_heading("2.3 Acteurs reellement presents", level=2)
    add_table(doc, ["Acteur", "Origine", "Description observee dans le code"], [
        ["USER", "Role realm Keycloak", "Utilisateur standard de la gateway (role realm de base)."],
        ["INTERN", "Role realm Keycloak", "Utilisateur interne Synapse (« Internal Synapse user »)."],
        ["EXTERN", "Role realm Keycloak", "Utilisateur externe Synapse (« External Synapse user »)."],
        ["ADMIN", "Role realm Keycloak", "Administrateur de la gateway ; seul role autorise sur tous les endpoints /api/admin/** (@PreAuthorize(\"hasRole('ADMIN')\"))."],
        ["Compte de service gateway-admin", "Client Keycloak confidentiel", "Compte machine-a-machine (client_credentials) utilise par le backend pour appeler l'API d'administration Keycloak (gestion des utilisateurs/roles)."],
    ], col_widths=[3.2, 3.5, 9.3])
    note(doc, "les 8 comptes de demonstration importes dans le realm (admin, admin1, admin2, "
              "extern1, extern2, intern1, intern2, user) sont des donnees d'environnement de "
              "developpement local, pas des acteurs fonctionnels distincts.")

    doc.add_heading("2.4 Fonctionnalites principales (verifiees dans le code)", level=2)
    bullets(doc, [
        "Connexion via Keycloak (Authorization Code, redirection geree par keycloak-js).",
        "Conversations : creation implicite au premier message, liste paginee filtrable, renommage, archivage/restauration, suppression definitive, changement de modele en cours de conversation.",
        "Envoi de messages avec reponse en streaming (SSE) token par token.",
        "Envoi de pieces jointes (jusqu'a 10 par message) : documents (pdf, docx, pptx, csv, xlsx, zip), images, fichiers texte/code — analysees par le DLP avant envoi au modele.",
        "Inspection d'une piece jointe : contenu original, elements sensibles detectes, version securisee (masquee), telechargement de la version masquee.",
        "Blocage ou masquage automatique des messages/fichiers contenant des donnees sensibles, avec message explicatif en francais cote frontend et possibilite de renvoyer une version masquee (message « masquer et renvoyer » pour les blocages de severite medium).",
        "Selection du modele LLM parmi un catalogue actif (OpenAI, Groq, Gemini, Mistral), filtre selon les restrictions applicables a l'utilisateur/role.",
        "Console d'administration : vue d'ensemble/metriques securite, gestion des regles DLP et mots bannis, gestion des fournisseurs/modeles LLM, gestion des utilisateurs et roles Keycloak, journal d'audit des actions admin, journal des incidents DLP.",
    ])

    doc.add_heading("2.5 Contraintes techniques observees", level=2)
    bullets(doc, [
        "Le schema de base de donnees est gere exclusivement par Flyway (spring.jpa.hibernate.ddl-auto=validate) : toute evolution de schema doit passer par une migration versionnee.",
        "Le comportement du service DLP est fail-closed : toute erreur d'analyse (timeout, service indisponible, extraction impossible) entraine un blocage plutot qu'un passage silencieux.",
        "Aucune cle API de fournisseur LLM n'est stockee ni exposee par le backend ou la base de donnees ; seules des references de variables d'environnement (api_key_env_var) sont persistees.",
        "Le backend et le frontend ne sont pas conteneurises dans docker-compose.yml (ils s'executent en local via Maven/npm) ; seuls PostgreSQL, LiteLLM, le service DLP et Keycloak (+ sa base et son provisioner) sont des services Docker.",
        "Une seule variable d'environnement racine (.env, non versionnee) centralise l'ensemble des secrets locaux.",
    ])

    doc.add_heading("2.6 Vue globale du fonctionnement", level=2)
    add_figure(doc, D / "architecture_globale.png", "Figure 1 — Vue globale des composants et de leurs interactions.")


# ---------------------------------------------------------------------------
# 3. Architecture generale
# ---------------------------------------------------------------------------

def add_section_03(doc):
    doc.add_heading("3. Architecture generale", level=1)

    doc.add_heading("3.1 Style architectural", level=2)
    p(doc, "Architecture multi-services orientee API : un frontend SPA (Single Page Application) "
           "sans routeur URL consomme une API REST/SSE exposee par un backend monolithique Spring "
           "Boot organise en couches (controller / service / repository / entity). Le backend "
           "delegue deux responsabilites a des services externes specialises : l'authentification "
           "a Keycloak (OIDC/OAuth2) et l'analyse de securite du contenu au service DLP (FastAPI). "
           "L'acces aux modeles de langage est indirect, via le proxy LiteLLM qui normalise les "
           "API des differents fournisseurs (format compatible OpenAI).")

    doc.add_heading("3.2 Architecture logique", level=2)
    bullets(doc, [
        "Couche presentation : application React (composants + hooks, sans etat global type Redux).",
        "Couche API/orchestration : controllers Spring MVC (REST + SSE) et services applicatifs.",
        "Couche securite transversale : validation JWT (resource server OAuth2), verification de role (@PreAuthorize), verification DLP systematique avant tout appel LLM.",
        "Couche persistance : Spring Data JPA sur PostgreSQL, schema gere par Flyway.",
        "Couche integration externe : clients WebClient (reactifs) vers DLP, LiteLLM et l'API d'administration Keycloak.",
    ])

    doc.add_heading("3.3 Architecture applicative", level=2)
    p(doc, "Le backend est un module Maven unique (backend/, Spring Boot 4.1.0, Java 17) melangeant "
           "Spring MVC (controllers HTTP/SSE classiques) et Spring WebFlux (WebClient reactif pour "
           "les appels sortants vers DLP/LiteLLM/Keycloak) — MVC pour l'exposition, WebFlux pour la "
           "consommation de services externes, comme indique dans le README du projet. Le frontend "
           "est une application Vite/React independante (frontend/), sans framework de routage, "
           "dont toute la navigation est geree par etat React local et persistee ponctuellement en "
           "localStorage/IndexedDB.")

    doc.add_heading("3.4 Architecture technique", level=2)
    add_table(doc, ["Couche", "Technologie", "Port par defaut"], [
        ["Frontend", "React 19 / Vite 8 (serveur de dev)", "5173"],
        ["Backend API", "Spring Boot 4.1.0 / Java 17", "8081 (chemin /api)"],
        ["Service DLP", "Python 3.11 / FastAPI", "8000"],
        ["Proxy LLM", "LiteLLM (image docker.litellm.ai/berriai/litellm:latest)", "4000"],
        ["Authentification", "Keycloak 26.7.0", "8080"],
        ["Base applicative", "PostgreSQL 16-alpine", "5433 -> 5432 (conteneur)"],
        ["Base Keycloak", "PostgreSQL 16-alpine (instance dediee)", "interne au reseau Docker"],
    ], col_widths=[3.5, 8.5, 4.0])

    doc.add_heading("3.5 Composants principaux et responsabilites", level=2)
    add_table(doc, ["Composant", "Responsabilite"], [
        ["Frontend React", "Interface utilisateur (chat, conversations, modeles, administration), gestion du flux SSE, integration Keycloak cote client."],
        ["Backend Spring Boot", "Point d'entree unique de l'API, authentification/autorisation, orchestration DLP -> LiteLLM, persistance des conversations/messages/pieces jointes, administration."],
        ["Service DLP (FastAPI)", "Extraction de texte, detection de donnees sensibles (regex + Presidio NER + mots bannis), scoring, decision ALLOW/MASK/BLOCK, masquage."],
        ["LiteLLM", "Normalisation des appels vers les fournisseurs LLM externes (format compatible OpenAI), gestion des cles API fournisseurs."],
        ["Keycloak", "Authentification OIDC des utilisateurs, gestion des realms/roles/clients, API d'administration consommee par le backend."],
        ["PostgreSQL (application)", "Persistance du catalogue de modeles, des conversations/messages/pieces jointes, des permissions et de l'audit."],
    ], col_widths=[4.5, 11.5])

    doc.add_heading("3.6 Communications entre composants", level=2)
    bullets(doc, [
        "Frontend -> Backend : HTTPS/HTTP REST + SSE, en-tete Authorization: Bearer <JWT Keycloak> sur chaque requete (rafraichi via keycloak.updateToken avant chaque appel).",
        "Frontend -> Keycloak : redirection navigateur pour la connexion (flux Authorization Code), rafraichissement de session en arriere-plan.",
        "Backend -> Service DLP : appels HTTP REST synchrones (WebClient), timeout de connexion 2s / lecture 10s par defaut, corps JSON ou multipart selon texte/pieces jointes.",
        "Backend -> LiteLLM : appel HTTP compatible OpenAI (/v1/chat/completions), avec ou sans streaming (Server-Sent Events) selon le cas d'usage.",
        "Backend -> Keycloak (Admin API) : authentification machine-a-machine (client_credentials, client gateway-admin), puis appels REST sur /admin/realms/{realm}/...",
        "LiteLLM -> Fournisseurs LLM externes : appels HTTP sortants authentifies par cle API (OpenAI, Groq, Gemini, Mistral).",
    ])

    doc.add_heading("3.7 Dependances", level=2)
    p(doc, "Le backend depend au demarrage de PostgreSQL (application) pour Flyway/JPA, et au moment "
           "des requetes de LiteLLM, du service DLP et de Keycloak (validation JWT via issuer-uri, "
           "et Admin API pour les fonctionnalites d'administration des utilisateurs). Le frontend "
           "depend du backend pour toute donnee applicative et de Keycloak pour l'authentification. "
           "Le service DLP est independant (aucune dependance sortante vers les autres services) "
           "mais depend de modeles NLP telecharges a la construction de son image Docker (spaCy "
           "en_core_web_sm / fr_core_news_sm) et de Tesseract OCR (langues eng/fra/ara).")

    p(doc, "Voir la figure 1 (section 2.6) pour le diagramme d'architecture globale correspondant "
           "exactement aux composants et communications decrits dans cette section.")


# ---------------------------------------------------------------------------
# 4. Architecture Frontend
# ---------------------------------------------------------------------------

def add_section_04(doc):
    doc.add_heading("4. Architecture Frontend", level=1)

    doc.add_heading("4.1 Framework et version", level=2)
    add_table(doc, ["Dependance", "Version"], [
        ["react / react-dom", "^19.2.7"],
        ["vite", "^8.1.1"],
        ["keycloak-js", "^26.2.4"],
        ["react-markdown / remark-gfm / rehype-sanitize", "^10.1.0 / ^4.0.1 / ^6.0.0"],
        ["react-syntax-highlighter", "^16.1.1"],
        ["mammoth (conversion .docx cote client)", "^1.12.0"],
        ["vitest (tests)", "^4.0.15"],
        ["eslint", "^10.6.0"],
    ], col_widths=[9.0, 7.0])
    note(doc, "aucune bibliotheque de routage (react-router), de gestion d'etat globale "
              "(Redux/Zustand/…) ni de client HTTP tiers (axios) n'est presente dans "
              "frontend/package.json — confirme par inspection directe du fichier.")

    doc.add_heading("4.2 Structure du projet", level=2)
    add_table(doc, ["Dossier", "Role"], [
        ["src/api/", "Couche fetch unique vers le backend (client.js + un module par ressource : conversations, models, admin, attachments)."],
        ["src/components/common/", "Composants UI generiques (ConfirmDialog, SelectDropdown, Toast, icones)."],
        ["src/features/chat/", "Fil de discussion, composer, streaming SSE, rendu Markdown/code, gestion des pieces jointes, messages DLP."],
        ["src/features/conversations/", "Liste/historique des conversations, recherche, archivage."],
        ["src/features/layout/", "Coquille applicative : Sidebar, AppLayout, gestion des menus/panneaux, glisser-deposer de fichiers."],
        ["src/features/models/", "Catalogue de modeles, selecteur, galerie."],
        ["src/features/admin/", "Console d'administration (AdminDashboard, composants et utilitaires associes)."],
        ["src/hooks/", "Hooks generiques transverses (localStorage, clic exterieur)."],
        ["src/utils/", "Utilitaires (decodage JWT, formatage d'erreurs, metadonnees modeles, cles de stockage)."],
        ["src/styles/", "Une feuille CSS par domaine (chat, composer, sidebar, admin, tokens, etc.)."],
        ["public/assets/", "Icones et images statiques (PNG) referencees par chemin direct."],
    ], col_widths=[4.5, 11.5])

    doc.add_heading("4.3 Pages et routage", level=2)
    p(doc, "Il n'existe pas de routeur URL (aucune dependance react-router, aucun composant "
           "BrowserRouter/Routes trouve dans le code). L'application est une SPA a URL unique : la "
           "navigation « chat / catalogue de modeles / administration » est geree par de l'etat "
           "React local (activeView dans useAppMenus, showAdminDashboard dans App.jsx, "
           "activeSection dans AdminDashboard.jsx persiste en localStorage). Il n'y a donc pas de "
           "protection d'acces par route : toute l'application est protegee globalement par la "
           "garde d'authentification Keycloak (voir 4.7), et la visibilite de la console admin est "
           "une condition d'affichage cote client (voir 4.14).")

    doc.add_heading("4.4 Composants principaux", level=2)
    bullets(doc, [
        "AppLayout.jsx / Sidebar.jsx — coquille generale, navigation, panneau d'inspection redimensionnable, glisser-deposer de fichiers.",
        "ChatThread.jsx / ChatMessage.jsx / ChatComposer.jsx — affichage des messages, saisie, gestion des pieces jointes.",
        "MarkdownContent.jsx / CodeBlock.jsx — rendu Markdown securise (rehype-sanitize) et coloration syntaxique du code.",
        "DocumentInspectorPanel.jsx — panneau d'inspection d'une piece jointe (original / elements detectes / version securisee).",
        "ConversationList.jsx / ConversationItem.jsx / SearchModal.jsx — historique et recherche des conversations.",
        "ModelSelector.jsx / ModelGallery.jsx / ModelCard.jsx — selection et exploration des modeles.",
        "AdminDashboard.jsx — console d'administration complete (une seule vue par section, sans sous-routage).",
    ])

    doc.add_heading("4.5 Services (couche API)", level=2)
    p(doc, "src/api/client.js centralise toutes les requetes : construction de l'URL de base "
           "(VITE_API_BASE_URL, defaut /api), rafraichissement du jeton Keycloak avant chaque appel "
           "(keycloak.updateToken(30)), ajout de l'en-tete Authorization, gestion uniforme des "
           "erreurs (classe ApiError) et des reponses 204/401. Quatre modules specialises "
           "consomment ce client : conversationsApi.js, modelsApi.js, adminApi.js, attachmentsApi.js.")

    doc.add_heading("4.6 Hooks", level=2)
    add_table(doc, ["Hook", "Role"], [
        ["useChatController", "Orchestration transverse chat + conversations + modeles (envoi, ouverture, changement de modele)."],
        ["useChatUi / useMessageStream / useAutoScroll", "Etat du fil de discussion, cycle de vie du flux SSE, defilement automatique."],
        ["useConversations / useConversationStatus", "CRUD des conversations, filtres, statut d'activite client uniquement (idle/generating/completed_unread)."],
        ["useModels", "Chargement du catalogue de modeles, modele selectionne, persistance localStorage."],
        ["useAppMenus", "Etat des menus/panneaux/vue active de la coquille applicative."],
        ["useLocalStorage / useOutsideClick", "Hooks generiques reutilisables."],
    ], col_widths=[4.5, 11.5])

    doc.add_heading("4.7 Gestion de l'etat", level=2)
    p(doc, "Aucune bibliotheque de gestion d'etat globale : un seul React Context "
           "(AuthContext, expose par AuthProvider.jsx, porte l'instance keycloak-js) et le reste de "
           "l'etat applicatif est compose via des hooks personnalises (voir 4.6). "
           "AdminDashboard.jsx concentre a lui seul une trentaine d'etats locaux (useState) pour "
           "l'ensemble de la console d'administration, sans decomposition en hooks dedies.")

    doc.add_heading("4.8 Communication avec le backend", level=2)
    p(doc, "Toutes les requetes passent par fetch natif (aucun axios). Les flux de streaming "
           "(reponse LLM) ne sont pas geres via l'API EventSource native (qui ne permet pas "
           "d'en-tetes personnalises ni de corps POST) mais via response.body.getReader() et un "
           "TextDecoder, avec un decodage manuel des trames SSE separees par \\n\\n "
           "(src/features/chat/utils/sse.js).")

    doc.add_heading("4.9 Gestion de l'authentification", level=2)
    bullets(doc, [
        "Bibliotheque : keycloak-js v26.2.4 (adaptateur officiel, pas de wrapper React tiers).",
        "Initialisation : keycloak.init({ onLoad: 'login-required', checkLoginIframe: false }) — tant que l'utilisateur n'est pas authentifie, AuthProvider.jsx ne rend rien ; il n'existe pas de page de connexion propre au frontend, la redirection se fait vers la page de connexion hebergee par Keycloak (theme personnalise « synapse »).",
        "Jeton gere en memoire par keycloak-js (non persiste explicitement par le code applicatif) ; rafraichissement a la demande avant chaque appel API (pattern « refresh-on-demand », pas de minuteur de fond).",
        "Roles -> UI : utils/authUtils.js decode manuellement le payload JWT (sans verification de signature, a but d'affichage uniquement) pour determiner hasAdminRole a partir de realm_access.roles.",
        "Deconnexion : keycloak.logout({redirectUri: window.location.origin}).",
    ])
    note(doc, "l'adaptateur keycloak-js active PKCE (S256) par defaut depuis la version 18 ; le "
              "code applicatif ne force pas explicitement pkceMethod et l'attribut PKCE n'est pas "
              "explicitement positionne sur le client synapse-client dans le realm importe — "
              "PKCE est donc actif via le comportement par defaut de la bibliotheque plutot que par "
              "une politique explicite du realm ou du code. Le client synapse-client a egalement "
              "directAccessGrantsEnabled=true dans le realm (flux mot de passe direct disponible en "
              "plus du flux par redirection), point a verifier si une politique plus stricte est "
              "souhaitee.")

    doc.add_heading("4.10 Gestion des conversations", level=2)
    p(doc, "Une conversation n'est creee cote serveur qu'au premier envoi de message (pas de "
           "creation a vide). Le frontend propose : liste paginee avec filtres modele/recherche/"
           "archive, renommage en ligne, archivage (suppression logique) et restauration, "
           "suppression definitive (avec confirmation), et changement de modele en cours de "
           "conversation (avec choix « nouvelle conversation » ou « continuer »).")

    doc.add_heading("4.11 Gestion du streaming", level=2)
    p(doc, "Le flux SSE est lu manuellement (voir 4.8) et interprete evenement par evenement : "
           "message (confirmation serveur du message optimiste), token (ajout progressif de texte, "
           "effet machine a ecrire), done (fin, statut TERMINE), error (notamment code "
           "DLP_BLOCKED, qui transforme le message utilisateur en carte d'alerte DLP dans le fil de "
           "discussion plutot que d'afficher une simple erreur generique).")

    doc.add_heading("4.12 Gestion des fichiers", level=2)
    bullets(doc, [
        "Depot par bouton (trombone) ou glisser-deposer sur toute la fenetre.",
        "Restrictions cote client (pre-verification UX, non garantes de securite) : 10 pieces jointes maximum par message, liste blanche d'extensions (documents, images, texte, code).",
        "Persistance locale des pieces jointes non envoyees dans IndexedDB (survit a un rechargement de page) via features/chat/utils/pendingAttachments.js.",
        "Envoi en multipart/form-data vers l'endpoint de streaming avec fichiers.",
        "Inspection post-envoi : contenu original, elements sensibles detectes, version masquee — telechargeable ou renvoyable comme message securise.",
    ])

    doc.add_heading("4.13 Interface administrateur", level=2)
    p(doc, "AdminDashboard.jsx expose six sections (vue d'ensemble, securite, modeles, "
           "utilisateurs, roles, audit) rendues dans la meme coquille que le chat (bascule "
           "d'affichage animee par document.startViewTransition, pas de sous-routage). La "
           "navigation admin est integree directement dans le composant Sidebar.jsx principal "
           "(le composant AdminSidebar de AdminComponents.jsx n'est plus importe nulle part — code "
           "mort probable, voir section 22).")

    doc.add_heading("4.14 Gestion des erreurs", level=2)
    p(doc, "Un systeme de toasts global (components/common/Toast.jsx) affiche les erreurs et succes "
           "(auto-effacement apres quelques secondes). utils/errors.js centralise le formatage des "
           "messages (detection d'echec reseau/CORS, messages specifiques pour les limites de "
           "pieces jointes, messages generiques pour l'indisponibilite du modele). Les blocages DLP "
           "sont affiches en ligne dans le fil de discussion, pas via un toast.")

    doc.add_heading("4.15 Fonctionnalites reellement implementees vs controles uniquement visuels", level=2)
    p(doc, "Les elements suivants sont des restrictions cote client (UX) sans garantie de securite "
           "si elles ne sont pas dupliquees cote serveur :")
    bullets(doc, [
        "Affichage/masquage du bouton « Administration » selon le decodage local (non verifie cryptographiquement) du role ADMIN dans le JWT.",
        "Restriction de la liste des roles selectionnables dans la console admin a ADMIN/INTERN/EXTERN.",
        "Filtrage des extensions de fichiers acceptees a l'upload (attribut accept + verification JS).",
        "Limite client de 10 pieces jointes par message (une limite serveur distincte existe egalement, remontee via le code d'erreur ATTACHMENT_LIMIT_EXCEEDED).",
    ])
    p(doc, "A l'inverse, la decision DLP (ALLOW/MASK/BLOCK) est entierement determinee cote "
           "serveur : le frontend ne fait qu'afficher le resultat renvoye par le backend, sans "
           "aucune logique de decision cote client — separation correcte pour ce mecanisme.")


# ---------------------------------------------------------------------------
# 5. Architecture Backend
# ---------------------------------------------------------------------------

def add_section_05(doc):
    doc.add_heading("5. Architecture Backend", level=1)

    doc.add_heading("5.1 Framework et version", level=2)
    p(doc, "Spring Boot 4.1.0 (parent Maven), Java 17. Starters principaux : "
           "spring-boot-starter-webmvc (controllers REST/SSE), spring-boot-starter-webflux "
           "(WebClient reactif pour les integrations sortantes), spring-boot-starter-data-jpa, "
           "spring-boot-starter-oauth2-resource-server (validation JWT), "
           "spring-boot-starter-validation, spring-boot-starter-actuator, "
           "spring-boot-starter-flyway + flyway-database-postgresql, pilote org.postgresql:postgresql.")

    doc.add_heading("5.2 Organisation des packages (com.example.backend)", level=2)
    add_table(doc, ["Package", "Role"], [
        ["config", "Securite (SecurityConfig), CORS (CorsConfig), donnees de demarrage dev (StartupDataFixer)."],
        ["controller", "Controleurs REST (chat, conversations, pieces jointes, modeles, sante, ensemble des endpoints /admin/**)."],
        ["dto", "Objets requete/reponse (records) exposes par l'API."],
        ["entity", "Entites JPA persistees."],
        ["enums", "Enums metier (statuts, roles de message)."],
        ["exceptions", "Exceptions personnalisees (DLP, contenu non autorise)."],
        ["integration.dlp", "Client HTTP + DTOs vers le service DLP externe."],
        ["integration.keycloak", "Client HTTP vers l'API d'administration Keycloak."],
        ["integration.litellm", "Client HTTP + DTO vers le proxy LiteLLM."],
        ["repository", "Interfaces Spring Data JPA."],
        ["service", "Logique metier (chat, conversations, DLP, pieces jointes, utilisateur courant)."],
    ], col_widths=[3.5, 12.5])

    doc.add_heading("5.3 Controllers", level=2)
    p(doc, "Voir la liste exhaustive des endpoints en section 6. Vue d'ensemble des controllers : "
           "ChatController, ConversationController, AttachmentController, ModelController, "
           "HealthController (endpoints applicatifs authentifies) ; AdminMetricsController, "
           "AdminModelController, AdminPermissionController, RolePermissionController, "
           "PatternController, AuditLogController, FilteredMessageController, "
           "KeycloakAdminController (endpoints d'administration, tous restreints au role ADMIN).")
    note(doc, "/api/health n'est pas un endpoint public malgre son nom : il n'a aucune annotation "
              "@PreAuthorize propre mais reste soumis a la regle globale anyRequest().authenticated() "
              "de SecurityConfig. Le seul endpoint reellement public est /actuator/health "
              "(Spring Boot Actuator, explicitement permitAll()).")

    doc.add_heading("5.4 Services", level=2)
    add_table(doc, ["Service", "Responsabilite"], [
        ["ChatService", "Chat one-shot non-streame (endpoint de compatibilite POST /api/chat) : verifie le modele, valide l'acces, applique le DLP, appelle LiteLLM."],
        ["ChatValidationService", "Verifie l'acces aux modeles (restrictions utilisateur/role, contournement ADMIN uniquement sur restrictions de role, jamais sur restrictions personnelles) et agrege les mots bannis applicables."],
        ["ConversationService", "Service central (~1050 lignes) : cycle de vie complet des conversations et messages, preparation/streaming SSE, persistance des messages bloques par le DLP, construction du contexte envoye au LLM (fenetre glissante de 20 messages par defaut)."],
        ["DlpService", "Traduit les reponses du client DLP en decisions gateway, applique la politique fail-closed, journalise les incidents."],
        ["CurrentUserService", "Resout l'entite Utilisateur locale a partir du JWT (creation automatique si absente), extrait keycloakId et roles realm ; fail-closed (401) si sub absent ou non-UUID."],
        ["AttachmentService", "Stockage disque des pieces jointes (protection anti path-traversal), persistance des metadonnees DLP, endpoints d'inspection/telechargement securise, suppression physique a la suppression definitive d'une conversation."],
        ["MessagePersistenceService", "Persiste l'etat final (succes/echec) d'un message assistant depuis les callbacks asynchrones du streaming LiteLLM."],
        ["FilteredMessageAuditWriter", "Ecrit les incidents DLP dans une transaction independante (REQUIRES_NEW) pour ne jamais perdre l'audit meme en cas d'echec de la requete appelante."],
        ["DemoUserProvider", "Fournit l'utilisateur demo-user fixe ; non branche dans les controllers observes (vestige probable, voir section 22)."],
    ], col_widths=[4.2, 11.8])

    doc.add_heading("5.5 Repositories", level=2)
    p(doc, "Interfaces Spring Data JPA standard, avec requetes derivees ou JPQL personnalisees : "
           "UtilisateurRepository, ConversationRepository (recherche paginee avec jointures), "
           "MessageRepository (ordre, statut, nettoyage des liens de reponse), "
           "AttachmentRepository (verification de propriete via jointure jusqu'a l'utilisateur), "
           "ModeleLlmRepository, FournisseurLlmRepository, AuditLogRepository et "
           "FilteredMessageRepository (Specification JPA pour filtres dynamiques), "
           "GlobalBannedWordRepository, UserBannedWordRepository, RoleBannedWordRepository, "
           "UserLlmRestrictionRepository, RoleLlmRestrictionRepository.")

    doc.add_heading("5.6 Entities (JPA)", level=2)
    p(doc, "Voir le detail complet des tables et colonnes en section 7 (Architecture des donnees). "
           "Entites : Utilisateur, Conversation, Message, Attachment, ModeleLlm, FournisseurLlm, "
           "AuditLog, FilteredMessage, GlobalBannedWord, UserBannedWord, RoleBannedWord, "
           "UserLlmRestriction, RoleLlmRestriction.")

    doc.add_heading("5.7 DTOs", level=2)
    bullets(doc, [
        "Chat : ChatRequest{model, message}, ChatResponse{model, answer}.",
        "Conversation : CreateConversationRequest, UpdateConversationRequest, ChangeConversationModelRequest, ConversationResponse, ConversationPageResponse, SendMessageRequest, MessageResponse (inclut les champs DLP et pieces jointes).",
        "Modeles : ModelDto{alias, displayName, description, logoUrl, providerCode, providerName, status}.",
        "Pieces jointes : AttachmentContentResponse, AttachmentInspectionResponse, AttachmentSecureResponse.",
        "Admin : records internes aux controllers (AdminProviderResponse, AdminModelResponse, etc.).",
    ])
    note(doc, "deux DTOs declares (AddBannedWordRequest, AddLlmRestrictionRequest, avec validation "
              "Bean Validation) ne sont utilises par aucun controller : AdminPermissionController "
              "manipule en realite des Map<String,String> brutes sans validation typee — "
              "refactorisation probablement inachevee (voir section 22).")

    doc.add_heading("5.8 Configurations", level=2)
    bullets(doc, [
        "SecurityConfig — resource server OAuth2/JWT, CSRF desactive, sessions STATELESS, /actuator/health public, tout le reste authenticated(); convertisseur d'autorites lisant realm_access.roles (pas resource_access) et prefixant chaque role en ROLE_<NOM_MAJUSCULE>.",
        "CorsConfig — origines autorisees limitees a http://localhost:5173 et http://127.0.0.1:5173 (propriete app.cors.allowed-origins), methodes GET/POST/PUT/PATCH/DELETE/OPTIONS.",
        "StartupDataFixer (@Profile(\"dev\")) — insere au demarrage l'utilisateur demo-user et un utilisateur de test a UUID fixe ; actif uniquement en profil dev.",
        "Aucune classe @ConfigurationProperties : toute la configuration passe par @Value injectes directement dans les constructeurs des clients d'integration.",
        "Deux @RestControllerAdvice distincts : ApiExceptionHandler (DLP, requetes invalides, limites de pieces jointes) et ChatExceptionHandler (ContentNotAllowedException — voir section 22, jamais levee dans le code actuel).",
        "Validation : Bean Validation standard (jakarta.validation, @NotBlank/@NotNull/@Size) via @Valid dans les controllers ; pas de Validator personnalise.",
    ])

    doc.add_heading("5.9 Services externes consommes", level=2)
    add_table(doc, ["Service", "Client backend", "Protocole"], [
        ["LiteLLM", "integration.litellm.LiteLlmService", "HTTP REST, POST /v1/chat/completions, avec ou sans stream=true (SSE)"],
        ["Service DLP", "integration.dlp.DlpClient / DlpPatternClient", "HTTP REST (JSON ou multipart), WebClient avec timeouts dedies"],
        ["Keycloak (validation JWT)", "spring-boot-starter-oauth2-resource-server", "OIDC discovery via issuer-uri"],
        ["Keycloak (Admin API)", "integration.keycloak.KeycloakAdminClient", "OAuth2 client_credentials + REST Admin API"],
    ], col_widths=[3.5, 6.0, 6.5])

    doc.add_heading("5.10 Flux metier principaux", level=2)
    p(doc, "Voir le detail complet en section 13 (Flux techniques principaux) et le diagramme de "
           "sequence associe.")


# ---------------------------------------------------------------------------
# 6. API REST et Streaming
# ---------------------------------------------------------------------------

def add_section_06(doc):
    doc.add_heading("6. API REST et Streaming", level=1)
    p(doc, "Toutes les routes ci-dessous sont prefixees par /api. Sauf mention contraire, chaque "
           "endpoint exige un JWT Keycloak valide (regle globale anyRequest().authenticated() de "
           "SecurityConfig) ; la colonne « Authentification » precise en plus les exigences de role "
           "specifiques (@PreAuthorize).")

    rows = [
        ["GET", "/actuator/health", "Spring Boot Actuator", "Publique (permitAll)", "Etat de sante technique du processus backend."],
        ["GET", "/api/health", "HealthController", "JWT (aucun role)", "Etat applicatif ({status, service}) — non public malgre son nom."],
        ["POST", "/api/chat", "ChatController", "JWT (aucun role)", "Chat one-shot non-streame (compatibilite) : verifie modele + acces, applique DLP, appelle LiteLLM."],
        ["POST", "/api/conversations", "ConversationController", "JWT", "Cree une conversation (modelAlias, title)."],
        ["GET", "/api/conversations", "ConversationController", "JWT", "Liste paginee, filtres modelAlias/search/archived/page/size."],
        ["GET", "/api/conversations/{id}", "ConversationController", "JWT (propriete verifiee)", "Detail d'une conversation."],
        ["PATCH", "/api/conversations/{id}", "ConversationController", "JWT (propriete)", "Renomme la conversation (title)."],
        ["PATCH", "/api/conversations/{id}/model", "ConversationController", "JWT (propriete)", "Change le modele associe (modelAlias)."],
        ["DELETE", "/api/conversations/{id}", "ConversationController", "JWT (propriete)", "Archive la conversation (suppression logique)."],
        ["PATCH", "/api/conversations/{id}/restore", "ConversationController", "JWT (propriete)", "Restaure une conversation archivee."],
        ["DELETE", "/api/conversations/{id}/permanent", "ConversationController", "JWT (propriete)", "Suppression definitive (supprime aussi les fichiers pieces jointes sur disque)."],
        ["GET", "/api/conversations/{id}/messages", "ConversationController", "JWT (propriete)", "Liste des messages de la conversation."],
        ["POST", "/api/conversations/{id}/messages/stream", "ConversationController", "JWT (propriete)", "SSE — envoie un message texte, reponse assistant en streaming (evenements message/token/done/error)."],
        ["POST", "/api/conversations/{id}/messages/stream-with-files", "ConversationController", "JWT (propriete)", "SSE, multipart/form-data — envoie texte + jusqu'a 10 pieces jointes."],
        ["GET", "/api/attachments/{id}", "AttachmentController", "JWT (propriete)", "Metadonnees d'une piece jointe."],
        ["GET", "/api/attachments/{id}/content", "AttachmentController", "JWT (propriete)", "Fichier original brut."],
        ["GET", "/api/attachments/{id}/inspection", "AttachmentController", "JWT (propriete)", "Texte extrait + correspondances DLP publiques."],
        ["GET", "/api/attachments/{id}/secure", "AttachmentController", "JWT (propriete)", "Texte masque de la piece jointe."],
        ["GET", "/api/attachments/{id}/secure/download", "AttachmentController", "JWT (propriete)", "Telechargement du texte masque (.txt)."],
        ["POST", "/api/conversations/{cid}/attachments/{id}/send-secure", "AttachmentController", "JWT (propriete)", "SSE — envoie le texte masque de la piece jointe comme message."],
        ["GET", "/api/models", "ModelController", "JWT", "Alias de modeles accessibles a l'utilisateur courant (filtres restrictions)."],
        ["GET", "/api/models/details", "ModelController", "JWT", "Detail enrichi des modeles (List<ModelDto>)."],
        ["GET", "/api/admin/metrics/security", "AdminMetricsController", "ADMIN", "Statistiques DLP (incidents totaux/bloques/masques, severites, evenements d'audit)."],
        ["GET", "/api/admin/models/providers", "AdminModelController", "ADMIN", "Liste des fournisseurs LLM."],
        ["POST", "/api/admin/models/providers", "AdminModelController", "ADMIN", "Cree un fournisseur."],
        ["PATCH", "/api/admin/models/providers/{id}", "AdminModelController", "ADMIN", "Modifie un fournisseur."],
        ["DELETE", "/api/admin/models/providers/{id}", "AdminModelController", "ADMIN", "Supprime un fournisseur (409 si modeles associes)."],
        ["PATCH", "/api/admin/models/providers/{id}/status", "AdminModelController", "ADMIN", "Active/desactive un fournisseur."],
        ["GET", "/api/admin/models", "AdminModelController", "ADMIN", "Liste des modeles."],
        ["POST", "/api/admin/models", "AdminModelController", "ADMIN", "Cree un modele."],
        ["PATCH", "/api/admin/models/{id}", "AdminModelController", "ADMIN", "Modifie un modele."],
        ["DELETE", "/api/admin/models/{id}", "AdminModelController", "ADMIN", "Supprime un modele (409 si reference)."],
        ["PATCH", "/api/admin/models/{id}/status", "AdminModelController", "ADMIN", "Active/desactive un modele."],
        ["POST", "/api/admin/models/{id}/test", "AdminModelController", "ADMIN", "Teste la connexion au modele via LiteLLM (latence, statut CONNECTED/FAILED)."],
        ["GET", "/api/admin/permissions/users", "AdminPermissionController", "ADMIN", "Liste tous les Utilisateur internes."],
        ["GET/POST/DELETE", "/api/admin/permissions/banned-words/global{,/{id}}", "AdminPermissionController", "ADMIN", "Mots bannis globaux."],
        ["GET/POST/DELETE", "/api/admin/permissions/llm-restrictions/{userKeycloakId}, /llm-restrictions{,/{id}}", "AdminPermissionController", "ADMIN", "Restrictions de modele par utilisateur."],
        ["GET/POST/DELETE", "/api/admin/permissions/banned-words/user/{userKeycloakId}, /user{,/{id}}", "AdminPermissionController", "ADMIN", "Mots bannis par utilisateur."],
        ["GET/POST/DELETE", "/api/admin/permissions/banned-words/role/{roleName}, /role{,/{id}}", "RolePermissionController", "ADMIN", "Mots bannis par role."],
        ["GET/POST/DELETE", "/api/admin/permissions/llm-restrictions/role/{roleName}, /role{,/{id}}", "RolePermissionController", "ADMIN", "Restrictions de modele par role."],
        ["GET", "/api/admin/permissions/patterns", "PatternController", "ADMIN", "Lit le fichier de regles DLP (patterns.json)."],
        ["POST", "/api/admin/permissions/patterns", "PatternController", "ADMIN", "Ajoute une regle DLP, synchronise vers le service DLP."],
        ["PUT", "/api/admin/permissions/patterns/{name}", "PatternController", "ADMIN", "Met a jour une regle DLP existante."],
        ["DELETE", "/api/admin/permissions/patterns/{name}", "PatternController", "ADMIN", "Supprime une regle DLP."],
        ["GET", "/api/admin/audit", "AuditLogController", "ADMIN", "Journal d'audit pagine/filtrable (search, action, entityName, performedBy, from, to ; taille max 100)."],
        ["GET", "/api/admin/filtered-messages", "FilteredMessageController", "ADMIN", "Journal des incidents DLP pagine/filtrable."],
        ["GET", "/api/admin/keycloak/users", "KeycloakAdminController", "ADMIN", "Liste des utilisateurs Keycloak (search), proxy Admin API."],
        ["POST", "/api/admin/keycloak/users", "KeycloakAdminController", "ADMIN", "Cree un utilisateur Keycloak (201)."],
        ["PATCH", "/api/admin/keycloak/users/{id}/enabled", "KeycloakAdminController", "ADMIN", "Active/desactive un utilisateur Keycloak."],
        ["GET", "/api/admin/keycloak/roles", "KeycloakAdminController", "ADMIN", "Roles realm geres (filtres selon keycloak.admin.managed-roles)."],
        ["GET", "/api/admin/keycloak/users/{id}/roles", "KeycloakAdminController", "ADMIN", "Roles realm de l'utilisateur."],
        ["PUT", "/api/admin/keycloak/users/{id}/roles", "KeycloakAdminController", "ADMIN", "Remplace les roles (exactement un role exige dans le payload)."],
    ]
    add_table(doc, ["Methode", "Endpoint", "Controller", "Authentification", "Description"], rows,
              col_widths=[2.0, 6.0, 3.0, 2.2, 8.0])

    doc.add_heading("6.1 Streaming (SSE)", level=2)
    p(doc, "Quatre endpoints exposent un flux Server-Sent Events (SseEmitter cote backend) : "
           "POST /api/conversations/{id}/messages/stream, "
           "POST /api/conversations/{id}/messages/stream-with-files, "
           "POST /api/conversations/{cid}/attachments/{id}/send-secure. "
           "Le flux emet des evenements nommes message, token, done et error — voir sections 8 et "
           "13 pour le detail du protocole et le diagramme de sequence.")

    doc.add_heading("6.2 Endpoints lies aux fichiers", level=2)
    p(doc, "Voir AttachmentController (section 6, tableau ci-dessus) pour le cycle de vie complet "
           "d'une piece jointe : metadonnees, contenu original, inspection DLP, version securisee, "
           "telechargement, et envoi de la version securisee comme message.")

    doc.add_heading("6.3 Endpoints administrateur", level=2)
    p(doc, "L'ensemble des routes /api/admin/** est protege par @PreAuthorize(\"hasRole('ADMIN')\") "
           "(au niveau classe pour la plupart des controllers, au niveau methode pour "
           "RolePermissionController et PatternController). Chaque mutation (creation, "
           "modification, suppression, changement de statut) declenche une ecriture dans la table "
           "audit_logs (action, entite, identifiant, performedBy issu du sub du JWT).")

    doc.add_heading("6.4 Endpoints de securite", level=2)
    p(doc, "La synchronisation des regles DLP (PatternController) et les metriques de securite "
           "(AdminMetricsController) sont reservees au role ADMIN. Cote service DLP lui-meme, deux "
           "endpoints d'administration (GET/PUT /admin/patterns) sont proteges par une cle "
           "d'administration distincte (en-tete X-DLP-Admin-Key) — voir section 10.")


# ---------------------------------------------------------------------------
# 7. Architecture des donnees
# ---------------------------------------------------------------------------

def add_section_07(doc):
    doc.add_heading("7. Architecture des donnees", level=1)

    doc.add_heading("7.1 SGBD et strategie de migration", level=2)
    p(doc, "PostgreSQL 16 (image Docker postgres:16-alpine, deux instances distinctes : une pour "
           "l'application, une dediee a Keycloak). Le schema applicatif est gere exclusivement par "
           "Flyway (spring.flyway.enabled=true, spring.flyway.locations=classpath:db/migration) ; "
           "Hibernate est configure en spring.jpa.hibernate.ddl-auto=validate, c'est-a-dire qu'il "
           "verifie la coherence entite <-> table au demarrage mais ne genere ni n'altere jamais le "
           "schema. 22 migrations versionnees (V1 a V22, numerotation continue et sans trou) sont "
           "appliquees dans backend/src/main/resources/db/migration.")
    note(doc, "un dossier backend/src/main/resources/db/migration_backup/ contient deux scripts "
              "(anciennes versions V10/V13) qui ne sont pas dans le chemin classpath:db/migration "
              "et ne sont donc jamais executes — vestiges d'une renumerotation passee, sans effet "
              "sur le schema reel.")

    doc.add_heading("7.2 Historique des migrations Flyway", level=2)
    migs = [
        ["V1", "Cree fournisseur_llm et modele_llm (catalogue LLM) + index."],
        ["V2", "Seed : 4 fournisseurs (openai, groq, gemini, mistral) et 4 modeles actifs."],
        ["V3", "Ajoute modele_llm.nom_affichage (rempli puis NOT NULL)."],
        ["V4", "Cree utilisateur, conversation, message + index ; insere l'utilisateur demo demo-user."],
        ["V5", "Ajoute message.modele_llm_id (FK), retro-remplit pour les messages ASSISTANT."],
        ["V6", "Recree les FK message.conversation_id (CASCADE) et reponse_a_message_id (SET NULL)."],
        ["V7", "Correctif de donnees ponctuel sur le modele Gemini (rendu obsolete par V8)."],
        ["V8", "Bascule le modele Gemini vers gemini/gemini-3.6-flash."],
        ["V9", "Elargit message.statut pour inclure DLP_BLOCKED ; ajoute dlp_highest_severity, dlp_detected_types."],
        ["V10", "Ajoute message.dlp_matches_json."],
        ["V11", "Ajoute message.dlp_masked_text."],
        ["V12", "Nettoyage retroactif : reecrit le contenu des messages DLP_BLOCKED persistes avant le durcissement du pipeline DLP."],
        ["V13", "Ajoute message.attachment_metadata_json."],
        ["V14", "Cree la table attachment (pieces jointes liees a un message)."],
        ["V15", "Cree 5 tables de permissions/restrictions (user_llm_restrictions, role_llm_restrictions, global_banned_words, user_banned_words, role_banned_words)."],
        ["V16", "Cree audit_logs."],
        ["V17", "Cree filtered_messages (journal DLP au niveau requete)."],
        ["V18", "Etend filtered_messages (highest_severity, detected_types, detection_count, request_status)."],
        ["V19", "Ajoute modele_llm.description et logo_url."],
        ["V20", "Ajoute fournisseur_llm.api_key_env_var (nom de variable d'env, jamais le secret)."],
        ["V21", "Change modele_llm.logo_url de VARCHAR(500) vers TEXT (logos encodes en donnees base64)."],
        ["V22", "Remplace le modele Groq deprecie par groq/openai/gpt-oss-20b."],
    ]
    add_table(doc, ["Version", "Resume"], migs, col_widths=[1.8, 14.2])

    doc.add_heading("7.3 Tables principales", level=2)
    p(doc, "Etat final apres application des 22 migrations :")
    tables = [
        ["fournisseur_llm", "id (PK), code (UNIQUE), nom, statut (CHECK ACTIF/INACTIF), api_key_env_var, created_at/updated_at"],
        ["modele_llm", "id (PK), fournisseur_llm_id (FK), alias_interne (UNIQUE), nom_modele_provider, nom_affichage, description, logo_url (TEXT), statut"],
        ["utilisateur", "id (PK), external_id (UNIQUE), nom_affichage, created_at/updated_at"],
        ["conversation", "id (PK), utilisateur_id (FK), modele_llm_id (FK), titre, statut (CHECK ACTIVE/ARCHIVEE), dernier_message_at"],
        ["message", "id (PK), conversation_id (FK CASCADE), modele_llm_id (FK, nullable), reponse_a_message_id (FK auto-reference, SET NULL), role, ordre, statut, contenu, colonnes DLP (dlp_highest_severity, dlp_detected_types, dlp_matches_json, dlp_masked_text), attachment_metadata_json"],
        ["attachment", "id (PK), message_id (FK CASCADE), original_filename, storage_key, mime_type, size, dlp_decision, extraction_status, extracted_text/masked_text/matches_json"],
        ["user_llm_restrictions / role_llm_restrictions", "Restrictions d'acces aux modeles par utilisateur (UUID Keycloak) ou par role (nom) — pas de FK vers modele_llm (reference par alias texte)"],
        ["global_banned_words / user_banned_words / role_banned_words", "Mots bannis globaux / par utilisateur / par role — pas de FK"],
        ["audit_logs", "id (PK), action, entity_name, entity_id, performed_by (UUID), timestamp — pas de FK"],
        ["filtered_messages", "id (PK), user_keycloak_id (UUID), original_content, redacted_content, action, reason, highest_severity, detected_types, detection_count, request_status, timestamp — pas de FK"],
    ]
    add_table(doc, ["Table", "Colonnes cles / contraintes"], tables, col_widths=[3.8, 12.2])

    doc.add_heading("7.4 Relations", level=2)
    bullets(doc, [
        "fournisseur_llm (1) -> modele_llm (N).",
        "utilisateur (1) -> conversation (N).",
        "modele_llm (1) -> conversation (N) [modele par defaut de la conversation].",
        "conversation (1) -> message (N), ON DELETE CASCADE.",
        "message (1) -> message (N), auto-reference reponse_a_message_id, ON DELETE SET NULL.",
        "modele_llm (1) -> message (N), attribution du modele ayant produit une reponse ASSISTANT (FK nullable).",
        "message (1) -> attachment (N), ON DELETE CASCADE.",
    ])
    note(doc, "les tables de permissions, d'audit et d'incidents DLP (user_llm_restrictions, "
              "role_llm_restrictions, global/user/role_banned_words, audit_logs, filtered_messages) "
              "referencent l'identite par UUID Keycloak ou par nom de role, sans contrainte de cle "
              "etrangere SQL vers utilisateur ou message — choix d'architecture assume puisque "
              "l'identite est geree par un realm Keycloak externe a PostgreSQL.")

    doc.add_heading("7.5 Contraintes et index", level=2)
    bullets(doc, [
        "Contrainte UNIQUE composite uq_message_conversation_ordre (conversation_id, ordre).",
        "Index composite idx_conversation_utilisateur_statut_dernier_message pour la liste paginee des conversations.",
        "Index idx_message_conversation_statut_role_ordre, idx_message_reponse_a, idx_message_modele_llm pour les acces frequents au fil de messages.",
        "Contraintes CHECK (plutot que des types ENUM PostgreSQL) sur fournisseur_llm.statut, modele_llm.statut, conversation.statut, message.role, message.statut.",
        "Aucune contrainte CHECK SQL sur attachment.dlp_decision/extraction_status ni sur filtered_messages.action/request_status — validation uniquement applicative (Java), point a noter.",
    ])

    doc.add_heading("7.6 Donnees seedees", level=2)
    p(doc, "V2 insere 4 fournisseurs et 4 modeles actifs (secure-gpt, secure-groq, secure-gemini, "
           "secure-mistral). V4 insere un unique utilisateur de demonstration (external_id="
           "'demo-user'), sans mot de passe ni secret associe. Aucune migration n'insere de secret, "
           "cle API ou mot de passe en clair ou hache.")

    add_figure(doc, D / "erd_donnees.png", "Figure 2 — Modele conceptuel de donnees (schema PostgreSQL, V1-V22).")


# ---------------------------------------------------------------------------
# 8. Authentification et autorisation
# ---------------------------------------------------------------------------

def add_section_08(doc):
    doc.add_heading("8. Authentification et autorisation", level=1)

    doc.add_heading("8.1 Keycloak — realm et clients", level=2)
    p(doc, "Le mecanisme d'authentification est exclusivement Keycloak/OIDC (realm synapse, "
           "Keycloak 26.7.0). Il n'existe pas d'authentification locale (pas de table de mots de "
           "passe applicative). Trois clients Keycloak sont pertinents pour ce systeme :")
    add_table(doc, ["Client", "Type", "Usage"], [
        ["synapse-client", "Public (Authorization Code, PKCE par defaut de keycloak-js)", "Client du frontend React ; redirectUris limites a localhost:5173/127.0.0.1:5173 ; directAccessGrantsEnabled=true egalement actif."],
        ["gateway-admin", "Confidentiel, service account (client_credentials)", "Utilise par le backend pour appeler l'API d'administration Keycloak ; roles limites a query-users, view-users, manage-users, view-realm sur realm-management."],
        ["security-admin-console / account(-console)", "Clients Keycloak integres", "Console d'administration Keycloak elle-meme (PKCE S256 explicitement configure sur ces clients build-in) ; non utilises par l'application."],
    ], col_widths=[3.5, 5.0, 7.5])

    doc.add_heading("8.2 Roles realm", level=2)
    p(doc, "Quatre roles realm applicatifs : ADMIN, INTERN, EXTERN, USER (confirmes dans le realm "
           "importe et dans le script de provisioning). Le role par defaut de Keycloak "
           "(default-roles-synapse) est compose des roles techniques standard (offline_access, "
           "uma_authorization, gestion du profil de compte).")

    doc.add_heading("8.3 Authentification frontend", level=2)
    p(doc, "keycloak-js initialise en onLoad: 'login-required' — l'application entiere est bloquee "
           "(rien n'est rendu) tant que l'utilisateur n'est pas authentifie ; la redirection se fait "
           "vers la page de connexion hebergee par Keycloak avec le theme personnalise « synapse » "
           "(login/theme.properties, CSS et logo dedies, localisation francaise). Le jeton est geree "
           "en memoire par la bibliotheque et rafraichi a la demande avant chaque appel API.")

    doc.add_heading("8.4 Authentification backend", level=2)
    p(doc, "Spring Security en mode resource server OAuth2 : chaque requete sur /api/** (hors "
           "/actuator/health) doit porter un JWT Bearer valide, verifie via "
           "spring.security.oauth2.resourceserver.jwt.issuer-uri "
           "(http://127.0.0.1:8080/realms/synapse par defaut). Le convertisseur d'autorites "
           "personnalise lit realm_access.roles (pas resource_access) et prefixe chaque role en "
           "ROLE_<NOM_MAJUSCULE> pour l'utiliser avec hasRole(...).")

    doc.add_heading("8.5 Protection des endpoints", level=2)
    bullets(doc, [
        "Regle globale SecurityConfig : anyRequest().authenticated(), sauf /actuator/health (permitAll).",
        "Tous les controllers sous /api/admin/** portent @PreAuthorize(\"hasRole('ADMIN')\") (au niveau classe pour la plupart, au niveau methode pour RolePermissionController et PatternController).",
        "Les endpoints de conversations/pieces jointes verifient en plus la propriete de la ressource par l'utilisateur courant (requetes repository findOwnedById filtrant par utilisateur), et pas seulement le role.",
        "CurrentUserService echoue en 401 (fail-closed) si le claim sub du JWT est absent ou n'est pas un UUID valide.",
    ])

    doc.add_heading("8.6 Controle d'acces cote serveur vs cote frontend", level=2)
    p(doc, "Le controle reellement applique se situe cote backend (validation cryptographique du "
           "JWT par Spring Security, verification du role via @PreAuthorize, verification de "
           "propriete par requete repository). Le frontend ne fait que refleter ces droits pour "
           "l'ergonomie : le decodage du role ADMIN dans utils/authUtils.js est un decodage base64 "
           "manuel sans verification de signature, utilise uniquement pour afficher ou masquer le "
           "bouton « Administration » — voir section 4.15 pour la liste complete des controles "
           "visuels uniquement.")

    doc.add_heading("8.7 Administration des utilisateurs et roles", level=2)
    p(doc, "La gestion des comptes et des roles ne passe pas par une table locale mais entierement "
           "par l'API d'administration Keycloak, via KeycloakAdminController (role ADMIN requis) et "
           "KeycloakAdminClient (compte de service gateway-admin). Les roles geres/exposes par "
           "cette API sont restreints a la liste keycloak.admin.managed-roles (par defaut "
           "ADMIN,INTERN,EXTERN) — empechant la modification de roles techniques internes de "
           "Keycloak.")


# ---------------------------------------------------------------------------
# 9. Architecture de securite
# ---------------------------------------------------------------------------

def add_section_09(doc):
    doc.add_heading("9. Architecture de securite", level=1)

    doc.add_heading("9.1 Vue d'ensemble", level=2)
    p(doc, "La securite du systeme repose sur quatre piliers verifies dans le code : (1) "
           "authentification/autorisation Keycloak-JWT pour chaque requete API, (2) analyse DLP "
           "systematique de tout contenu (texte et pieces jointes) avant envoi a un LLM externe, "
           "(3) restrictions d'acces configurables par utilisateur/role sur les modeles et les mots "
           "bannis, (4) journalisation d'audit des actions d'administration et des incidents DLP.")

    doc.add_heading("9.2 DLP — moteur et integration", level=2)
    p(doc, "Le service DLP (dlp-service, FastAPI) combine un moteur de regles regex maison (39 "
           "regles declaratives dans patterns.json) et Microsoft Presidio (presidio-analyzer, "
           "moteur NLP spaCy fr/en) pour la detection d'entites nommees, complete par des "
           "recognizers Presidio personnalises (identifiants marocains : CIN, RIB, IBAN, telephone ; "
           "secrets techniques : cle OpenAI/AWS, token GitHub, JWT, cle privee, mot de passe en "
           "dur, chaine de connexion base de donnees). Le masquage (anonymisation) est realise par "
           "un module maison, pas par presidio-anonymizer (absent des dependances).")

    doc.add_heading("9.3 Classification et severite", level=2)
    add_table(doc, ["Severite", "Types de donnees (extraits du code, policy.py)"], [
        ["high", "carte de credit, IBAN, CIN marocaine, cle API/OpenAI, token GitHub, JWT, bearer token, cle privee, secret/mot de passe en dur, chaine de connexion DB, compte bancaire, BIC/SWIFT, portefeuille crypto, CVV, numero d'etat civil, passeport, IMEI, adresse IP"],
        ["medium", "identifiant alphanumerique, email, telephone, nom de personne"],
        ["low", "date de naissance, localisation, URL, organisation"],
    ], col_widths=[2.5, 13.5])

    doc.add_heading("9.4 Decisions ALLOW / MASK / BLOCK", level=2)
    p(doc, "Regle de decision (dlp/app/pipeline/decision.py) : la presence d'au moins une "
           "correspondance de severite high entraine BLOCK ; a defaut, la presence de toute "
           "correspondance (medium ou low) entraine MASK ; en l'absence de correspondance, ALLOW. "
           "En BLOCK, le message n'est jamais transmis a LiteLLM. En ALLOW/MASK, c'est toujours le "
           "texte renvoye par le DLP (masked_text) qui est transmis au LLM — jamais le texte "
           "original brut, meme en ALLOW (le DLP renvoie alors un texte identique a l'original).")

    doc.add_heading("9.5 Fail-closed", level=2)
    p(doc, "Le comportement fail-closed est explicite et documente a deux niveaux : cote service "
           "DLP, toute erreur d'extraction ou d'analyse renvoie une decision BLOCK plutot que de "
           "laisser passer un contenu non analyse (README du service : « Fail closed is "
           "intentional »). Cote backend, la Javadoc de DlpClient indique explicitement que les "
           "erreurs de transport (timeout, erreur HTTP, JSON invalide) sont volontairement "
           "converties en DlpUnavailableException afin que les couches superieures « echouent "
           "fermees » — aucun message n'atteint alors LiteLLM ; le client recoit une erreur "
           "HTTP 503 / evenement SSE error avec code DLP_UNAVAILABLE.")
    note(doc, "l'endpoint POST /analyse du service DLP (texte JSON pur) ne comporte pas de bloc "
              "try/except autour de l'appel au pipeline dans main.py : une exception interne "
              "inattendue y remonterait sous forme d'erreur HTTP 500 generique de FastAPI plutot "
              "que du format BLOCK habituel du service. Cote backend, cette reponse 500 est malgre "
              "tout interceptee comme une erreur de transport et convertie en fail-closed "
              "(DlpUnavailableException), donc le comportement observable de bout en bout reste "
              "fail-closed, mais la coherence interne du service DLP sur ce cas precis merite d'etre "
              "durcie.")

    doc.add_heading("9.6 Timeouts et gestion des erreurs (integration backend)", level=2)
    add_table(doc, ["Parametre", "Valeur par defaut"], [
        ["dlp.connect-timeout", "2 s"],
        ["dlp.read-timeout", "10 s"],
        ["litellm timeout (chat et streaming)", "60 s"],
        ["keycloak.admin timeout (Admin API)", "10 s"],
    ], col_widths=[7.0, 9.0])

    doc.add_heading("9.7 Controle des fichiers", level=2)
    p(doc, "Toute piece jointe est analysee par le DLP avant d'etre utilisee dans un message envoye "
           "au LLM (voir sections 10 et 12). Cote backend, le stockage disque applique une "
           "protection anti path-traversal (resolveStorageKey verifie que le chemin resolu reste "
           "sous le repertoire racine configure) et une limite de 10 pieces jointes par message.")

    doc.add_heading("9.8 Secrets et donnees sensibles", level=2)
    bullets(doc, [
        "Aucune cle API de fournisseur LLM n'est stockee en base de donnees ; seule une reference au nom de la variable d'environnement (fournisseur_llm.api_key_env_var) est persistee (V20).",
        "Toutes les cles/API secretes sont injectees via le fichier .env racine (non versionne, exclu par .gitignore) et consommees directement par LiteLLM/le backend/Keycloak.",
        "Le fichier .env.example ne contient que des valeurs de substitution (change_me_*, your_*_api_key_here, sk-local-litellm, dev-*).",
        "Le secret du client Keycloak gateway-admin est parametre (${GATEWAY_ADMIN_CLIENT_SECRET}) et non code en dur dans le realm importe.",
    ])

    doc.add_heading("9.9 Audit", level=2)
    p(doc, "Table audit_logs (action, entity_name, entity_id, performed_by, timestamp) ecrite a "
           "chaque mutation effectuee via un endpoint d'administration (creation/modification/"
           "suppression de fournisseur, modele, restriction, mot banni, regle DLP, utilisateur "
           "Keycloak). Table filtered_messages ecrite independamment (transaction REQUIRES_NEW, "
           "voir FilteredMessageAuditWriter) a chaque incident DLP (blocage ou masquage) sur le "
           "flux de texte simple.")


# ---------------------------------------------------------------------------
# 10. Architecture DLP
# ---------------------------------------------------------------------------

def add_section_10(doc):
    doc.add_heading("10. Architecture DLP", level=1)

    doc.add_heading("10.1 Technologies", level=2)
    add_table(doc, ["Composant", "Technologie / version"], [
        ["Framework web", "FastAPI 0.116.1, Uvicorn 0.35.0 (Python 3.11-slim)"],
        ["Validation", "Pydantic 2.11.7"],
        ["Detection NER", "presidio-analyzer 2.2.359, spaCy 3.7.5 (modeles en_core_web_sm, fr_core_news_sm)"],
        ["Detection de langue", "langdetect 1.0.9 (+ heuristique Unicode pour l'arabe)"],
        ["Extraction documents", "PyMuPDF 1.26.3 (pdf), python-docx 1.2.0, python-pptx 1.0.2, openpyxl 3.1.5"],
        ["OCR", "pytesseract 0.3.13 + Tesseract OCR (langues eng, fra, ara), pillow 11.3.0"],
        ["Upload", "python-multipart 0.0.20"],
        ["Tests", "pytest 8.4.1, httpx 0.28.1"],
    ], col_widths=[4.5, 11.5])
    note(doc, "presidio-anonymizer n'est pas une dependance du service : le masquage est realise "
              "par un module maison (dlp/app/pipeline/masking.py), pas par Presidio. Une couche "
              "NER par modele transformer existe dans le code (transformer_detector.py) mais est "
              "desactivee par defaut (DLP_TRANSFORMER_ENABLED) et sa dependance n'est pas installee "
              "par defaut.")

    doc.add_heading("10.2 Architecture du service (dlp/app)", level=2)
    add_table(doc, ["Module", "Role"], [
        ["main.py", "Point d'entree FastAPI, definition des endpoints, orchestration du pipeline."],
        ["config.py", "Lecture centralisee des variables d'environnement DLP_*."],
        ["policy.py", "Table statique de severite par type de donnee (TYPE_SEVERITY)."],
        ["schemas.py", "Modeles Pydantic des requetes/reponses."],
        ["detectors/", "regex_detector.py (+ patterns.json), presidio_detector.py, presidio_config.py, moroccan_recognizers.py, technical_secret_recognizers.py, transformer_detector.py, banned_words.py, language.py, luhn.py, iban.py."],
        ["ingestion/", "Parsers par format (pdf, docx, pptx, csv, xlsx, zip), OCR, routage des colonnes structurees, extension autorisees."],
        ["pipeline/", "normalization.py, evidence.py (score de confiance), dedup.py, ids.py, masking.py, decision.py, alerting.py."],
    ], col_widths=[3.5, 12.5])

    doc.add_heading("10.3 Endpoints", level=2)
    add_table(doc, ["Methode", "Endpoint", "But"], [
        ["POST", "/analyse", "Analyse un texte brut (JSON)."],
        ["POST", "/analyse-image", "OCR + analyse d'une image."],
        ["POST", "/analyse-pdf", "Extraction texte + OCR PDF, puis analyse."],
        ["POST", "/analyse-file", "Dispatch selon extension (docx/pptx/csv/xlsx/zip/texte brut)."],
        ["POST", "/analyse-message", "Analyse combinee texte + N pieces jointes ; decision agregee (BLOCK si une source BLOCK, sinon MASK si une source MASK, sinon ALLOW)."],
        ["GET", "/health", "Liveness."],
        ["GET", "/ready", "Readiness (force le chargement des modeles NLP ; 503 si echec)."],
        ["GET", "/admin/patterns", "Lit les regles regex actives (protege par X-Dlp-Admin-Key)."],
        ["PUT", "/admin/patterns", "Remplace integralement les regles regex (protege)."],
    ], col_widths=[2.0, 4.0, 10.0])

    doc.add_heading("10.4 Categories detectees", level=2)
    p(doc, "39 regles regex declaratives (patterns.json) et des recognizers Presidio predefinis "
           "(EMAIL_ADDRESS, PHONE_NUMBER, CREDIT_CARD, IBAN_CODE, IP_ADDRESS, URL, ORGANIZATION, "
           "PERSON) et personnalises (identifiants marocains, secrets techniques — voir 9.3). Les "
           "entites LOCATION sont volontairement toujours ignorees (choix produit assume dans le "
           "code et le README du service).")

    doc.add_heading("10.5 Pipeline de traitement", level=2)
    add_figure(doc, D / "pipeline_dlp.png", "Figure 3 — Pipeline de traitement du service DLP.")
    p(doc, "Etapes : controle de longueur (DLP_MAX_TEXT_LENGTH, defaut 50 000 caracteres) -> "
           "detection de langue -> normalisation (NFKC, anti-obfuscation d'e-mails) -> detection "
           "multi-source combinee (regex + Presidio + transformer optionnel + mots bannis fournis "
           "par la requete) -> suppression des correspondances deja neutralisees -> deduplication "
           "(priorite aux regles explicites) -> attribution d'identifiants lisibles (email_1, "
           "email_2, …) -> filtrage des correspondances configurees ALLOW -> decision -> "
           "journalisation des correspondances high (alerting) -> masquage (remplacement par des "
           "placeholders du type [EMAIL_1]).")

    doc.add_heading("10.6 Ingestion de fichiers", level=2)
    bullets(doc, [
        "Formats geres avec parseur dedie : .docx .pptx .csv .xlsx .pdf .zip ; formats texte brut et code source ; images (OCR).",
        "Limite de taille appliquee a deux niveaux : middleware sur Content-Length (contournable en transfert chunked, limite documentee dans le code) puis relecture bornee du flux (second filet).",
        "Limite de 10 pieces jointes par requete /analyse-message (DLP_MAX_ATTACHMENTS).",
        "Traitement ZIP securise (zip_parser.py) : profondeur maximale (DLP_MAX_ZIP_DEPTH, 3), nombre de fichiers maximal (DLP_MAX_ZIP_FILES, 50), taille decompressee totale maximale (DLP_MAX_ZIP_UNCOMPRESSED_MB, 50 Mo), ratio de compression maximal 100x (protection zip-bomb non pilotee par variable d'environnement), rejet des entrees chiffrees, protection anti path-traversal.",
    ])

    doc.add_heading("10.7 Comportement BLOCK / MASK / ALLOW", level=2)
    p(doc, "Voir 9.4 pour la regle de decision. Le champ matches de la reponse ne contient jamais la "
           "valeur sensible detectee elle-meme (strip_sensitive_values retire la cle value) — "
           "seules les metadonnees (type, position, severite, source) sont exposees.")

    doc.add_heading("10.8 Traitement des erreurs et timeouts", level=2)
    p(doc, "Toute erreur d'extraction (fichier corrompu, OCR en echec) renvoie une reponse "
           "structuree HTTP 200 avec status=ERROR, decision=BLOCK — un choix fail-closed explicite. "
           "Aucun mecanisme de timeout n'est configure cote service DLP lui-meme (la gestion du "
           "delai est deleguee a l'appelant, c'est-a-dire au backend — voir 9.6).")

    doc.add_heading("10.9 Interaction avec le backend", level=2)
    p(doc, "Le backend (DlpClient, DlpService) appelle /analyse ou /analyse-message selon le cas "
           "(texte seul ou avec pieces jointes), et PUT /admin/patterns (avec l'en-tete "
           "X-DLP-Admin-Key) pour synchroniser toute modification de regle faite via la console "
           "d'administration. Voir section 13 pour le diagramme de sequence complet.")

    doc.add_heading("10.10 Evaluation de la precision", level=2)
    p(doc, "dlp/evaluation/ contient un harnais distinct des tests unitaires : un corpus etiquete "
           "de 46 echantillons (FR/EN/AR, positifs/negatifs/adversariaux) et un script (evaluate.py) "
           "calculant precision/rappel/F1. Ces rapports (baseline-regex.json, current-regex.json) "
           "sont des artefacts versionnes de comparaison avant/apres modification des regles, pas "
           "une mesure en continu du service en production.")


# ---------------------------------------------------------------------------
# 11. Integration LLM
# ---------------------------------------------------------------------------

def add_section_11(doc):
    doc.add_heading("11. Integration LLM", level=1)

    doc.add_heading("11.1 LiteLLM — configuration", level=2)
    p(doc, "litellm/config.yaml definit le catalogue de modeles exposes par le proxy (image "
           "docker.litellm.ai/berriai/litellm:latest, port 4000). Aucune configuration de routeur "
           "ou de fallback multi-modele n'est definie : chaque alias pointe vers un seul provider.")
    add_table(doc, ["Alias interne", "Modele fournisseur", "Variable d'environnement requise"], [
        ["secure-gpt", "openai/gpt-4o-mini", "OPENAI_API_KEY"],
        ["secure-groq", "groq/openai/gpt-oss-20b", "GROQ_API_KEY"],
        ["secure-gemini", "gemini/gemini-3.6-flash", "GEMINI_API_KEY"],
        ["secure-mistral", "mistral/mistral-small-latest", "MISTRAL_API_KEY"],
        ["secure-claude (desactive, commente)", "anthropic/claude-3-5-sonnet-20241022", "ANTHROPIC_API_KEY"],
    ], col_widths=[4.5, 6.0, 5.0])
    note(doc, "litellm/README.md documente encore secure-groq -> groq/llama-3.1-8b-instant, un "
              "alias obsolete : la configuration reelle (config.yaml) est la source de verite et "
              "utilise groq/openai/gpt-oss-20b depuis la migration V22.")

    doc.add_heading("11.2 Parametres generaux LiteLLM", level=2)
    bullets(doc, [
        "general_settings.master_key = os.environ/LITELLM_MASTER_KEY (jamais en dur).",
        "litellm_settings : drop_params=true, set_verbose=true, request_timeout=60.",
        "Aucun hook DLP/guardrail natif LiteLLM configure : le controle DLP est entierement realise en amont, cote backend, avant tout appel a LiteLLM — LiteLLM n'a aucune connaissance du DLP.",
    ])

    doc.add_heading("11.3 Appel backend -> LiteLLM", level=2)
    p(doc, "integration.litellm.LiteLlmService construit une requete compatible OpenAI "
           "(POST /v1/chat/completions) avec l'en-tete Authorization: Bearer <LITELLM_MASTER_KEY="
           "********>. Deux modes : chat() (non-streame, timeout 60 s, extrait choices[0].message."
           "content) et streamChat() (stream=true, Accept: text/event-stream, parsing manuel des "
           "trames data: {...}). Chaque requete est prefixee d'un message systeme fixe (regles en "
           "francais) rappelant a l'assistant de ne jamais reconstruire les placeholders DLP "
           "([EMAIL_1], etc.) — ce message systeme n'est jamais persiste en base.")

    doc.add_heading("11.4 Streaming", level=2)
    p(doc, "Le flux de tokens SSE de LiteLLM est relaye au frontend par le backend sous forme de "
           "ses propres evenements SSE (token, done, error) — voir section 13 pour le detail "
           "complet du protocole.")

    doc.add_heading("11.5 Traitement des reponses et erreurs", level=2)
    bullets(doc, [
        "Reponse vide ou choices[0] absent -> ResponseStatusException(BAD_GATEWAY).",
        "Pour les modeles dont le nom contient « gemini » (insensible a la casse), safety_settings est force a BLOCK_NONE pour harcelement et contenu dangereux.",
        "Tokens vides filtres ; trame [DONE] ignoree explicitement.",
    ])

    doc.add_heading("11.6 Interaction avec le DLP", level=2)
    p(doc, "Aucun texte n'atteint LiteLLM sans etre d'abord passe par le service DLP : en cas de "
           "decision ALLOW ou MASK, c'est le texte renvoye par le DLP (masked_text) qui est "
           "transmis, jamais le texte utilisateur brut ; en cas de BLOCK, LiteLLM n'est jamais "
           "appele. Voir section 9.4 et le diagramme de sequence de la section 13.")


# ---------------------------------------------------------------------------
# 12. Gestion des fichiers
# ---------------------------------------------------------------------------

def add_section_12(doc):
    doc.add_heading("12. Gestion des fichiers", level=1)

    doc.add_heading("12.1 Upload", level=2)
    p(doc, "Depot par bouton ou glisser-deposer cote frontend ; envoi en multipart/form-data vers "
           "POST /api/conversations/{id}/messages/stream-with-files. Limite serveur : 10 pieces "
           "jointes par message (ConversationService.MAX_ATTACHMENTS_PER_MESSAGE), taille de "
           "fichier max 10 Mo et taille de requete max 25 Mo par defaut "
           "(spring.servlet.multipart.max-file-size / max-request-size, configurables).")

    doc.add_heading("12.2 Validation et formats supportes", level=2)
    p(doc, "Le service DLP definit la liste de reference des formats geres avec un parseur dedie "
           "(docx, pptx, csv, xlsx, pdf, zip) et des formats texte/code/image (voir section 10.6). "
           "Le frontend applique une liste blanche d'extensions similaire a titre de pre-verification "
           "UX uniquement (voir 4.15) ; l'application effective des restrictions revient au backend "
           "et au service DLP.")

    doc.add_heading("12.3 Stockage", level=2)
    p(doc, "Stockage sur disque local (app.attachments.storage-dir, defaut storage/attachments), "
           "confirme sur le disque du depot en backend/storage/attachments/{conversationId}/"
           "{messageId}/. Cle de stockage : {conversationId}/{messageId}/{UUID}-{filename}, avec "
           "verification que le chemin resolu reste sous le repertoire racine configure "
           "(protection anti path-traversal).")

    doc.add_heading("12.4 Analyse DLP", level=2)
    p(doc, "Chaque fichier est envoye au service DLP via /analyse-message (avec le texte du message "
           "eventuel) avant tout envoi au LLM. Le texte extrait, le texte masque et les "
           "correspondances (serialises en JSON dans une colonne TEXT) sont persistes pour permettre "
           "l'inspection ulterieure sans re-televerser le fichier.")

    doc.add_heading("12.5 Recuperation, previsualisation et telechargement", level=2)
    bullets(doc, [
        "GET /api/attachments/{id}/content — fichier original brut.",
        "GET /api/attachments/{id}/inspection — texte extrait + correspondances DLP publiques (sans valeur sensible).",
        "GET /api/attachments/{id}/secure — texte masque.",
        "GET /api/attachments/{id}/secure/download — telechargement du texte masque en .txt.",
        "Cote frontend, DocumentInspectorPanel.jsx previsualise le document original (conversion .docx via mammoth cote client, coloration syntaxique pour le texte/code, rendu natif image/iframe sinon).",
    ])

    doc.add_heading("12.6 Relation avec les conversations", level=2)
    p(doc, "Une piece jointe est toujours rattachee a un message (attachment.message_id, FK "
           "ON DELETE CASCADE) — donc indirectement a une conversation. La suppression definitive "
           "d'une conversation supprime egalement les fichiers physiques associes sur disque "
           "(AttachmentService.deleteFilesForConversation).")

    doc.add_heading("12.7 Controle d'acces", level=2)
    p(doc, "AttachmentRepository.findOwnedById verifie par jointure (message -> conversation -> "
           "utilisateur) que la piece jointe appartient bien a l'utilisateur authentifie avant de la "
           "servir — controle d'acces applique cote serveur pour chaque endpoint de "
           "AttachmentController.")

    doc.add_heading("12.8 Envoi securise", level=2)
    p(doc, "POST /api/conversations/{cid}/attachments/{id}/send-secure permet de renvoyer la "
           "version masquee d'une piece jointe deja analysee comme un nouveau message (flux SSE "
           "dedie, meme protocole que l'envoi de message standard).")


# ---------------------------------------------------------------------------
# 13. Flux techniques principaux
# ---------------------------------------------------------------------------

def add_section_13(doc):
    doc.add_heading("13. Flux techniques principaux", level=1)

    doc.add_heading("13.1 Authentification", level=2)
    p(doc, "L'utilisateur ouvre le frontend -> keycloak-js (onLoad: 'login-required') redirige vers "
           "la page de connexion Keycloak (theme synapse) -> apres connexion, redirection vers le "
           "frontend avec un jeton de session -> le frontend joint ce jeton (Authorization: Bearer) "
           "a chaque appel API, en le rafraichissant a la demande (updateToken(30)) -> le backend "
           "valide le JWT via l'issuer OIDC configure et en extrait les roles realm.")

    doc.add_heading("13.2 Creation d'une conversation", level=2)
    p(doc, "Aucun appel explicite de creation a vide : POST /api/conversations n'est declenche par "
           "le frontend qu'au moment du tout premier envoi de message (ensureConversation() cote "
           "frontend), avec un titre derive du contenu du premier message.")

    doc.add_heading("13.3 Envoi d'un message et controle DLP (ALLOW / MASK / BLOCK)", level=2)
    add_figure(doc, D / "sequence_chat_dlp.png", "Figure 4 — Sequence d'envoi d'un message avec controle DLP et streaming.")
    p(doc, "Le frontend appelle POST /api/conversations/{id}/messages/stream (ou "
           "…/stream-with-files si des pieces jointes sont presentes). ConversationService "
           "appelle DlpService, qui appelle le service DLP (/analyse ou /analyse-message). Selon la "
           "decision : BLOCK -> le message est persiste avec statut DLP_BLOCKED, journalise dans "
           "filtered_messages, et un evenement SSE error (code DLP_BLOCKED) est renvoye — LiteLLM "
           "n'est jamais appele. ALLOW/MASK -> le texte (original ou masque) est envoye a LiteLLM en "
           "mode streaming ; les tokens recus sont relayes au frontend via des evenements SSE token, "
           "puis un evenement done cloture le flux et declenche la persistance finale du message "
           "assistant (MessagePersistenceService).")

    doc.add_heading("13.4 Chargement de l'historique", level=2)
    p(doc, "GET /api/conversations (liste paginee/filtrable) puis GET /api/conversations/{id}/"
           "messages pour le detail d'une conversation selectionnee — chargement a la demande, pas "
           "de mise en cache serveur particuliere au-dela de la base de donnees elle-meme.")

    doc.add_heading("13.5 Ajout et analyse d'une piece jointe", level=2)
    p(doc, "Le fichier est joint au message (drag-and-drop ou bouton), persiste temporairement en "
           "IndexedDB cote client, puis envoye avec le texte via …/stream-with-files. Cote backend, "
           "chaque fichier est soumis au meme pipeline DLP que le texte (via /analyse-message), et "
           "son resultat (texte extrait, masque, correspondances) est persiste en base "
           "(table attachment) pour une inspection posterieure sans re-analyse.")

    doc.add_heading("13.6 Action administrateur", level=2)
    p(doc, "Toute mutation effectuee par un utilisateur ADMIN via la console d'administration "
           "(creation/modification/suppression de fournisseur ou modele, ajout/retrait de "
           "restriction ou de mot banni, modification d'une regle DLP, creation/modification d'un "
           "utilisateur ou de ses roles Keycloak) declenche : (1) l'operation elle-meme via l'API "
           "correspondante, (2) une ecriture dans audit_logs (action, entite, identifiant, "
           "performedBy issu du sub du JWT), et pour les regles DLP, (3) une synchronisation "
           "immediate vers le service DLP (PUT /admin/patterns).")


# ---------------------------------------------------------------------------
# 14. Docker et infrastructure
# ---------------------------------------------------------------------------

def add_section_14(doc):
    doc.add_heading("14. Docker et infrastructure", level=1)

    doc.add_heading("14.1 Services Docker Compose", level=2)
    add_table(doc, ["Service", "Image / Build", "Port (hote -> conteneur)", "Dependances", "Role"], [
        ["postgres", "postgres:16-alpine", "5433 -> 5432", "aucune", "Base de donnees applicative."],
        ["litellm", "docker.litellm.ai/berriai/litellm:latest", "4000 -> 4000", "aucune (appelle les fournisseurs externes)", "Proxy unifie vers les fournisseurs LLM."],
        ["dlp-service", "Build local (./dlp/Dockerfile)", "127.0.0.1:8000 -> 8000", "aucune", "Moteur DLP (FastAPI)."],
        ["keycloak-db", "postgres:16-alpine", "interne uniquement", "aucune", "Base de donnees dediee a Keycloak."],
        ["keycloak", "quay.io/keycloak/keycloak:26.7.0", "8080 -> 8080", "keycloak-db (healthy)", "Serveur d'authentification OIDC (import du realm synapse)."],
        ["keycloak-provisioner", "quay.io/keycloak/keycloak:26.7.0", "aucun", "keycloak (started)", "Job one-shot idempotent : theme/langue, roles INTERN/EXTERN, mots de passe demo, client gateway-admin, roles du compte de service."],
    ], col_widths=[2.6, 4.2, 3.0, 3.0, 3.2])
    note(doc, "le backend (Spring Boot) et le frontend (React/Vite) ne sont pas conteneurises dans "
              "docker-compose.yml — ils s'executent en local via mvnw et npm run dev, confirme par "
              "l'absence de backend/Dockerfile et frontend/Dockerfile dans le depot (seul "
              "dlp/Dockerfile existe).")

    doc.add_heading("14.2 dlp/Dockerfile", level=2)
    p(doc, "Image de base python:3.11-slim ; installation via apt de curl et tesseract-ocr "
           "(+ langues eng/fra/ara) ; installation des dependances Python (requirements.txt) ; "
           "telechargement des modeles spaCy en_core_web_sm et fr_core_news_sm au moment du build ; "
           "execution en utilisateur non-root dlpuser ; HEALTHCHECK sur GET /health (intervalle "
           "30 s, timeout 5 s, 5 tentatives, periode de demarrage 60 s) ; commande de demarrage "
           "uvicorn app.main:app --host 0.0.0.0 --port 8000.")

    doc.add_heading("14.3 Healthchecks", level=2)
    bullets(doc, [
        "postgres / keycloak-db : pg_isready sur la base configuree.",
        "dlp-service : requete Python vers http://localhost:8000/health.",
        "Aucun healthcheck Docker natif configure pour litellm ou keycloak dans docker-compose.yml (Keycloak expose neanmoins KC_HEALTH_ENABLED=true en interne).",
    ])

    doc.add_heading("14.4 Volumes et reseaux", level=2)
    p(doc, "Deux volumes nommes persistants : secure_llm_postgres_data (base applicative) et "
           "keycloak_postgres_data (base Keycloak). Volumes montes en lecture seule pour la "
           "configuration : ./litellm/config.yaml, ./keycloak/import, ./keycloak/themes, "
           "./keycloak/provision.sh. Aucun reseau Docker personnalise defini explicitement (reseau "
           "par defaut du projet Compose).")

    doc.add_heading("14.5 Ordre de demarrage", level=2)
    p(doc, "keycloak-db doit etre sain (condition: service_healthy) avant le demarrage de keycloak ; "
           "keycloak-provisioner attend que keycloak soit demarre (condition: service_started) avant "
           "de s'executer, puis se termine (restart: \"no\"). Les autres services (postgres, "
           "litellm, dlp-service) demarrent independamment.")

    doc.add_heading("14.6 Variables d'environnement (infrastructure)", level=2)
    p(doc, "Voir le detail complet en section 15 (Configuration).")


# ---------------------------------------------------------------------------
# 15. Configuration
# ---------------------------------------------------------------------------

def add_section_15(doc):
    doc.add_heading("15. Configuration", level=1)

    doc.add_heading("15.1 Backend — application.properties", level=2)
    p(doc, "Fichier unique backend/src/main/resources/application.properties (pas de profil "
           "separe application-{profil}.properties). Import optionnel du .env racine via "
           "spring.config.import.")
    add_table(doc, ["Propriete", "Valeur / defaut"], [
        ["server.port", "8081"],
        ["litellm.base-url", "http://localhost:4000"],
        ["litellm.master-key", "${LITELLM_MASTER_KEY} — ******** (obligatoire, pas de defaut)"],
        ["gateway.context.max-messages", "20"],
        ["dlp.base-url", "http://localhost:8000"],
        ["dlp.connect-timeout / dlp.read-timeout", "2s / 10s"],
        ["dlp.admin-key", "******** (defaut dev-dlp-admin)"],
        ["dlp.patterns.file", "../dlp/app/detectors/patterns.json"],
        ["keycloak.admin.base-url / realm / token-realm / client-id", "http://127.0.0.1:8080 / synapse / synapse / gateway-admin"],
        ["keycloak.admin.client-secret", "******** (${GATEWAY_ADMIN_CLIENT_SECRET})"],
        ["keycloak.admin.managed-roles", "ADMIN,INTERN,EXTERN"],
        ["spring.datasource.url / username / password", "jdbc:postgresql://localhost:5433/secure_llm_gateway / secure_llm_user / ******** (defaut dev change_me_local_only)"],
        ["spring.flyway.enabled / locations", "true / classpath:db/migration"],
        ["spring.jpa.hibernate.ddl-auto", "validate"],
        ["spring.security.oauth2.resourceserver.jwt.issuer-uri", "http://127.0.0.1:8080/realms/synapse"],
        ["app.cors.allowed-origins", "http://localhost:5173,http://127.0.0.1:5173"],
        ["app.attachments.storage-dir", "storage/attachments"],
        ["spring.servlet.multipart.max-file-size / max-request-size", "10MB / 25MB"],
    ], col_widths=[6.0, 10.0])
    note(doc, "app.attachments.max-llm-characters (defaut 40000, utilise par DlpService pour "
              "limiter la taille du contexte envoye au LLM) n'apparait pas dans application."
              "properties : c'est une valeur par defaut Java (@Value) non surchargeable sans "
              "l'ajouter explicitement au fichier.")

    doc.add_heading("15.2 Frontend — variables Vite", level=2)
    add_table(doc, ["Variable", "Role", "Defaut applicatif si absente"], [
        ["VITE_API_BASE_URL", "URL de base de l'API backend", "/api (suppose un reverse proxy en production)"],
        ["VITE_KEYCLOAK_URL", "URL du serveur Keycloak", "http://127.0.0.1:8080"],
        ["VITE_KEYCLOAK_REALM", "Realm Keycloak", "synapse"],
        ["VITE_KEYCLOAK_CLIENT_ID", "Client public Keycloak", "synapse-client"],
    ], col_widths=[4.5, 7.0, 4.5])
    note(doc, "ces valeurs (URLs et client-id public) sont injectees dans le bundle client et ne "
              "sont pas des secrets — confirme par AUDIT_MERGE_INTEGRATION.md.")

    doc.add_heading("15.3 LiteLLM — litellm/config.yaml", level=2)
    p(doc, "Voir le detail complet en section 11.1. Cle maitresse chargee via "
           "os.environ/LITELLM_MASTER_KEY (jamais en dur dans le fichier).")

    doc.add_heading("15.4 Service DLP — variables d'environnement", level=2)
    add_table(doc, ["Variable", "Defaut (docker-compose.yml)"], [
        ["DLP_MAX_TEXT_LENGTH", "50000"],
        ["DLP_MAX_FILE_SIZE_MB", "20"],
        ["DLP_MAX_ATTACHMENTS", "10"],
        ["DLP_MAX_ZIP_UNCOMPRESSED_MB", "50"],
        ["DLP_MAX_ZIP_FILES", "50"],
        ["DLP_MAX_ZIP_DEPTH", "3"],
        ["DLP_LOG_LEVEL", "INFO"],
        ["DLP_ADMIN_KEY", "******** (defaut dev-dlp-admin, a definir en environnement sensible)"],
    ], col_widths=[6.0, 10.0])

    doc.add_heading("15.5 Keycloak", level=2)
    p(doc, "Realm synapse importe depuis keycloak/import/synapse-realm.json au premier demarrage "
           "(--import-realm, ignore si le realm existe deja). Variables sensibles substituees "
           "depuis .env au moment de l'import : KEYCLOAK_ADMIN_USERNAME/PASSWORD, "
           "KEYCLOAK_DEMO_PASSWORD, GATEWAY_ADMIN_CLIENT_SECRET (aucune valeur reelle n'est "
           "committee dans le realm JSON). Theme de connexion personnalise « synapse » monte en "
           "lecture seule (keycloak/themes).")

    doc.add_heading("15.6 PostgreSQL", level=2)
    p(doc, "Deux instances distinctes (application et Keycloak), chacune configuree par ses propres "
           "variables POSTGRES_DB/USER/PASSWORD (application) et KEYCLOAK_DB_NAME/USER/PASSWORD "
           "(Keycloak), toutes injectees depuis le .env racine.")

    doc.add_heading("15.7 .env.example (racine)", level=2)
    p(doc, "34 lignes, uniquement des valeurs de substitution. Couvre : cles API des 5 fournisseurs "
           "LLM, LITELLM_MASTER_KEY/PORT, DLP_ADMIN_KEY, configuration PostgreSQL applicative, "
           "configuration datasource Spring, configuration base/admin/demo Keycloak, "
           "GATEWAY_ADMIN_CLIENT_SECRET, KEYCLOAK_ADMIN_BASE_URL/REALM/TOKEN_REALM/MANAGED_ROLES, "
           "KEYCLOAK_ISSUER_URI, KEYCLOAK_HOST_PORT. Aucune valeur reelle presente.")


# ---------------------------------------------------------------------------
# 16. Logs, audit et tracabilite
# ---------------------------------------------------------------------------

def add_section_16(doc):
    doc.add_heading("16. Logs, audit et tracabilite", level=1)

    doc.add_heading("16.1 Audit des actions d'administration", level=2)
    p(doc, "Table audit_logs (action, entity_name, entity_id, performed_by, timestamp), ecrite a "
           "chaque mutation effectuee via un endpoint /api/admin/**. Consultable via "
           "GET /api/admin/audit (pagination, filtres search/action/entityName/performedBy/from/to, "
           "taille de page plafonnee a 100), et depuis la section « Journal d'audit » de la console "
           "d'administration.")

    doc.add_heading("16.2 Journal des incidents DLP", level=2)
    p(doc, "Table filtered_messages (user_keycloak_id, original_content, redacted_content, action, "
           "reason, highest_severity, detected_types, detection_count, request_status, timestamp), "
           "ecrite dans une transaction independante (REQUIRES_NEW) par "
           "FilteredMessageAuditWriter afin de ne jamais perdre l'incident meme si la requete "
           "appelante echoue ensuite. Consultable via GET /api/admin/filtered-messages.")
    note(doc, "d'apres le code lu, cette ecriture d'audit DLP est realisee pour le flux de texte "
              "simple (DlpService.safeUserMessage) ; elle n'a pas ete confirmee comme systematique "
              "pour le flux pieces jointes (safeMessageForLlm) lors de la revue du code.")

    doc.add_heading("16.3 Logs applicatifs DLP", level=2)
    p(doc, "Le service DLP journalise en niveau WARNING (logger dlp_alerts, controle par "
           "DLP_LOG_LEVEL) chaque correspondance de severite high sous forme de JSON structure "
           "(timestamp, request_id genere par requete, user_id, decision, filename, liste des "
           "correspondances high) — uniquement les metadonnees, jamais la valeur sensible brute.")

    doc.add_heading("16.4 Metriques de securite", level=2)
    p(doc, "GET /api/admin/metrics/security agrege les statistiques DLP (incidents totaux, "
           "bloques, masques, severites critique/haute, compteurs du jour, nombre d'evenements "
           "d'audit) affichees dans la vue d'ensemble de la console d'administration.")

    doc.add_heading("16.5 Donnees stockees vs non stockees", level=2)
    bullets(doc, [
        "Stocke : contenu original et masque des messages/pieces jointes bloques ou masques, types de donnees detectes, severite, identifiant Keycloak de l'auteur.",
        "Jamais stocke : la valeur sensible individuelle dans les correspondances exposees par l'API (matches) ; les cles API des fournisseurs LLM.",
    ])


# ---------------------------------------------------------------------------
# 17. Gestion des erreurs et resilience
# ---------------------------------------------------------------------------

def add_section_17(doc):
    doc.add_heading("17. Gestion des erreurs et resilience", level=1)

    doc.add_heading("17.1 Gestionnaires d'exceptions (backend)", level=2)
    add_table(doc, ["Exception", "HTTP", "Code API"], [
        ["DlpBlockedException", "422 (UNPROCESSABLE_CONTENT)", "DLP_BLOCKED"],
        ["DlpUnavailableException (+ DlpInvalidResponseException)", "503", "DLP_UNAVAILABLE"],
        ["ResponseStatusException", "porte par l'exception (defaut 400)", "REQUEST_INVALID"],
        ["AttachmentLimitExceededException", "400", "ATTACHMENT_LIMIT_EXCEEDED"],
        ["ContentNotAllowedException", "422 (UNPROCESSABLE_ENTITY)", "content_not_allowed"],
    ], col_widths=[6.5, 5.0, 4.5])
    note(doc, "ContentNotAllowedException et son gestionnaire ChatExceptionHandler existent dans "
              "le code mais aucune partie du backend ne leve actuellement cette exception — "
              "mecanisme defini mais non cable (voir section 22).")

    doc.add_heading("17.2 Comportement quand le DLP est indisponible", level=2)
    p(doc, "Fail-closed strict : toute erreur de transport vers le service DLP (timeout, erreur "
           "HTTP, reponse JSON invalide ou incomplete) est convertie en DlpUnavailableException. "
           "Aucun message n'est alors envoye a LiteLLM. Le client recoit HTTP 503 ou, en streaming, "
           "un evenement SSE error avec le code DLP_UNAVAILABLE, affiche par le frontend comme "
           "« Controle de securite indisponible. Votre message n'a pas ete envoye au modele. ». "
           "Voir aussi section 9.5 pour le point d'attention concernant l'endpoint /analyse du "
           "service DLP lui-meme.")

    doc.add_heading("17.3 Comportement quand LiteLLM est indisponible", level=2)
    p(doc, "L'appel HTTP vers LiteLLM (chat ou streamChat) est borne par un timeout de 60 secondes. "
           "Une reponse vide ou l'absence de choices[0] declenche ResponseStatusException"
           "(BAD_GATEWAY). Aucun mecanisme de nouvelle tentative (retry) automatique n'a ete "
           "identifie dans le code du backend pour cet appel.")

    doc.add_heading("17.4 Erreurs PostgreSQL", level=2)
    p(doc, "Aucune gestion applicative dediee des erreurs PostgreSQL n'a ete identifiee au-dela du "
           "comportement par defaut de Spring Data JPA (propagation d'exception, remontee en erreur "
           "HTTP generique via la gestion Spring standard si non interceptee par un "
           "@RestControllerAdvice specifique). Le schema est valide au demarrage (Flyway + "
           "Hibernate ddl-auto=validate) : un desalignement de schema empeche le demarrage de "
           "l'application plutot que de provoquer une erreur silencieuse a l'execution.")

    doc.add_heading("17.5 Streaming interrompu", level=2)
    p(doc, "Cote frontend, l'arret volontaire d'une generation (bouton stop) declenche un "
           "AbortController ; le contenu partiel deja recu est conserve s'il est non vide, sans "
           "afficher d'erreur (distinction explicite entre AbortError volontaire et echec reel dans "
           "useMessageStream.js). Cote backend, MessagePersistenceService persiste l'etat final "
           "(succes/echec) du message assistant dans une transaction dediee, execute depuis les "
           "callbacks asynchrones du flux LiteLLM (donc apres le retour du controller HTTP).")

    doc.add_heading("17.6 Strategie fail-closed / fail-open reellement utilisee", level=2)
    p(doc, "Le systeme applique une strategie fail-closed sur l'ensemble de la chaine de securite : "
           "DLP indisponible ou en erreur -> blocage (jamais de passage silencieux au LLM) ; JWT "
           "absent, invalide ou sub non conforme -> 401 ; role insuffisant -> 403. Aucun mecanisme "
           "de fail-open (contournement automatique de la securite en cas de panne) n'a ete "
           "identifie dans le code.")


# ---------------------------------------------------------------------------
# 18. Tests
# ---------------------------------------------------------------------------

def add_section_18(doc):
    doc.add_heading("18. Tests", level=1)
    p(doc, "Aucun rapport de couverture de code n'a ete trouve dans le depot (pas de jacoco.xml, "
           "coverage.xml, lcov.info ni configuration de couverture dans les fichiers de build) — "
           "aucun pourcentage de couverture n'est donc avance dans cette section, conformement au "
           "principe de ne documenter que ce qui est verifiable.")

    doc.add_heading("18.1 Tests backend (Java, JUnit 5 / Mockito)", level=2)
    p(doc, "18 classes de test sous backend/src/test, environ 91 methodes @Test et 9 "
           "@ParameterizedTest. Ce sont exclusivement des tests unitaires (Mockito, "
           "@ExtendWith(MockitoExtension.class)) ou des tests de controller avec MockMvc en mode "
           "standalone (sans contexte Spring complet) : aucune classe n'utilise @SpringBootTest, "
           "@WebMvcTest, @DataJpaTest ni Testcontainers.")
    bullets(doc, [
        "Services : ConversationServiceTest (le plus fourni), DlpServiceTest, ChatServiceTest, CurrentUserServiceTest, AttachmentServiceTest, ChatValidationServiceTest, MessagePersistenceServiceTest, FilteredMessageAuditWriterTest.",
        "Integrations : DlpClientTest (serveur HTTP local factice simulant le service DLP, teste le fail-closed sur erreur 500/JSON invalide/decision inconnue/timeout), LiteLlmServiceTest.",
        "Controllers : ChatControllerTest, ConversationControllerTest, AdminPermissionControllerTest, AuditLogControllerTest, FilteredMessageControllerTest, AdminModelLogoValidationTest, PatternControllerPathTest, ApiExceptionHandlerTest.",
    ])

    doc.add_heading("18.2 Tests frontend (Vitest)", level=2)
    p(doc, "18 fichiers *.test.js(x), environ 73 blocs it()/test(). Les composants sont testes via "
           "react-dom/server (renderToStaticMarkup) et des assertions sur le HTML produit, pas via "
           "un DOM interactif jsdom (aucun environnement jsdom configure dans vite.config.js). "
           "Couvre notamment : client API, formatage d'erreurs, decodage JWT (authUtils), rendu des "
           "messages de chat (ChatMessage.test.jsx, le plus fourni), parsing SSE, gestion des "
           "pieces jointes en attente, AdminDashboard, ModelLogo, Toast.")

    doc.add_heading("18.3 Tests DLP (pytest)", level=2)
    p(doc, "18 fichiers de test sous dlp/tests, environ 231 fonctions test_*. "
           "dlp/tests/test_main.py (34 tests) exerce les endpoints FastAPI via TestClient (en "
           "process, sans reseau). dlp/tests/detectors/test_regex_detector.py est le plus volumineux "
           "(87 tests, couvre la majorite des regles regex). Sont egalement couverts : Presidio "
           "(fakes), IBAN/Luhn/langue, parseurs d'ingestion (docx/pptx/csv/xlsx/zip — securite ZIP "
           "incluse), et le pipeline (dedup, masquage, normalisation, decision, IDs).")

    doc.add_heading("18.4 Tests d'integration / bout-en-bout", level=2)
    p(doc, "Aucun test trouve qui appelle un vrai endpoint HTTP a travers plusieurs services "
           "reellement demarres simultanement. Les deux tests qui s'en approchent sont en realite "
           "isoles : DlpClientTest demarre un serveur HTTP local factice (pas le vrai service "
           "Python), et test_main.py utilise un client de test FastAPI in-process.")

    doc.add_heading("18.5 Tests specifiques a la securite", level=2)
    bullets(doc, [
        "DLP block/mask/allow et fail-closed : DlpServiceTest, DlpClientTest, dlp/tests/pipeline/test_decision.py, test_masking.py, test_main.py.",
        "Controle d'acces par role / restrictions LLM : ChatValidationServiceTest (contournement ADMIN sur restriction de role mais jamais sur restriction personnelle).",
        "Permissions admin : AdminPermissionControllerTest.",
        "Validite JWT : CurrentUserServiceTest (rejet des sujets non-UUID/manquants).",
        "Role ADMIN frontend : authUtils.test.js.",
        "Non-exposition d'informations sensibles dans les erreurs : dlpErrors.test.js.",
        "Validation anti-injection des URLs de logo : AdminModelLogoValidationTest.",
    ])

    doc.add_heading("18.6 Integration continue", level=2)
    p(doc, "Un unique workflow GitHub Actions (.github/workflows/ci.yml), declenche sur push "
           "(branches main, integration-review-remediation, file-integration-merge) et sur toute "
           "pull request, avec 4 jobs paralleles :")
    add_table(doc, ["Job", "Contenu"], [
        ["backend", "Java 17 (Temurin), ./mvnw -B test"],
        ["frontend", "Node 22, npm ci -> npm run lint -> npm test -- --run -> npm run build"],
        ["dlp", "Python 3.11, installation Tesseract + modeles spaCy, python -m pytest -q"],
        ["compose", "docker compose config avec des valeurs factices (aucun secret reel dans le CI)"],
    ], col_widths=[3.0, 13.0])
    p(doc, "Aucun test n'est marque desactive/skip/xfail dans les trois suites (recherche "
           "exhaustive sans resultat).")


# ---------------------------------------------------------------------------
# 19. Deploiement et execution
# ---------------------------------------------------------------------------

def add_section_19(doc):
    doc.add_heading("19. Deploiement et execution", level=1)
    p(doc, "Procedure verifiee dans README.md, applicable a l'environnement de developpement local "
           "actuel du depot (aucune procedure de deploiement production n'est definie dans le "
           "depot).")

    doc.add_heading("19.1 Prerequis", level=2)
    bullets(doc, [
        "Git, Docker Desktop.",
        "Java 17 (backend).",
        "Node.js et npm (frontend).",
        "Au moins une cle API d'un fournisseur LLM pour tester un modele reel.",
    ])

    doc.add_heading("19.2 Configuration initiale", level=2)
    p(doc, "Copier .env.example vers .env a la racine du depot et renseigner les valeurs locales "
           "(cles API, mots de passe locaux, LITELLM_MASTER_KEY identique cote LiteLLM et backend).")

    doc.add_heading("19.3 Demarrage des services Docker", level=2)
    p(doc, "docker compose up -d postgres litellm keycloak-db keycloak keycloak-provisioner "
           "(le service dlp-service peut etre demarre separement ou avec ce meme groupe selon les "
           "besoins). Keycloak est disponible sur http://localhost:8080 apres import du realm "
           "synapse.")

    doc.add_heading("19.4 Backend", level=2)
    p(doc, "Definir les variables d'environnement locales (LITELLM_MASTER_KEY, "
           "SPRING_DATASOURCE_URL/USERNAME/PASSWORD, KEYCLOAK_ADMIN_*, GATEWAY_ADMIN_CLIENT_SECRET), "
           "puis cd backend && mvnw.cmd spring-boot:run. Flyway s'execute automatiquement au "
           "demarrage. Le backend importe egalement automatiquement le .env racine si lance depuis "
           "la racine du depot ou depuis backend/.")

    doc.add_heading("19.5 Frontend", level=2)
    p(doc, "cd frontend && npm install && npm run dev — disponible sur http://localhost:5173, "
           "appelant l'API sur http://127.0.0.1:8081/api (proxy Vite en developpement).")

    doc.add_heading("19.6 Service DLP", level=2)
    p(doc, "Demarre via Docker Compose (dlp-service) ; peut aussi etre execute localement avec "
           "uvicorn app.main:app --reload depuis dlp/ pour le developpement (necessite les "
           "dependances Python et Tesseract OCR installes localement).")

    doc.add_heading("19.7 Verification rapide", level=2)
    bullets(doc, [
        "Sante backend : GET http://127.0.0.1:8081/api/health (authentifie) ou GET .../actuator/health (public).",
        "Catalogue de modeles : GET http://127.0.0.1:8081/api/models/details.",
        "Test direct LiteLLM : POST http://localhost:4000/v1/chat/completions avec Authorization: Bearer <LITELLM_MASTER_KEY>.",
        "Sante DLP : GET http://localhost:8000/health et /ready.",
    ])

    doc.add_heading("19.8 Reinitialisation locale", level=2)
    p(doc, "docker compose down -v puis docker compose up -d postgres litellm supprime et recree "
           "le volume PostgreSQL applicatif (Flyway rejoue toutes les migrations au redemarrage du "
           "backend). Une procedure separee et destructive existe pour reinitialiser uniquement le "
           "volume Keycloak (voir README.md, section dediee) — a n'executer que sur un "
           "environnement local sans donnees a conserver.")


# ---------------------------------------------------------------------------
# 20. Arborescence du projet
# ---------------------------------------------------------------------------

def add_section_20(doc):
    doc.add_heading("20. Arborescence du projet", level=1)
    p(doc, "Arborescence simplifiee (dossiers principaux uniquement) — exclut node_modules, "
           "target, dist, .venv, caches et fichiers temporaires.")

    tree = """gatewayLLM/
├── backend/                  Backend Spring Boot (API REST + SSE)
│   ├── src/main/java/com/example/backend/   controllers, services, repositories, entities, DTO, config
│   ├── src/main/resources/db/migration/     migrations Flyway (V1 a V22)
│   ├── src/test/                            tests unitaires et de controller (JUnit 5 / Mockito)
│   └── storage/attachments/                 fichiers de pieces jointes stockes sur disque
├── frontend/                 Frontend React / Vite (SPA sans routeur)
│   ├── src/api/                             couche fetch vers le backend
│   ├── src/features/                        chat, conversations, layout, models, admin
│   ├── src/components/common/               composants UI generiques
│   └── public/assets/                       icones et images statiques
├── dlp/                      Service DLP Python / FastAPI
│   ├── app/detectors/                       regex, Presidio, recognizers marocains/secrets
│   ├── app/ingestion/                       parsers de documents (pdf, docx, pptx, csv, xlsx, zip) + OCR
│   ├── app/pipeline/                        normalisation, dedup, masquage, decision
│   ├── tests/                               tests pytest (detectors, ingestion, pipeline)
│   └── evaluation/                          harnais d'evaluation precision/rappel (corpus etiquete)
├── database/                 Documentation de la base (README seul, pas de script SQL)
├── keycloak/
│   ├── import/                              realm synapse-realm.json
│   ├── themes/synapse/                      theme de connexion personnalise
│   └── provision.sh                         script de provisioning idempotent
├── litellm/
│   ├── config.yaml                          catalogue de modeles et parametres du proxy
│   └── examples/                            exemples de requetes de test
├── docs/                     Documentation complementaire (diagrammes PlantUML/PNG existants, notes de phase)
├── documents/                 Dossier d'Architecture Technique (ce document et son generateur)
├── .github/workflows/ci.yml  Integration continue (backend, frontend, dlp, compose)
├── docker-compose.yml         Orchestration des services d'infrastructure
├── .env.example                Modele de variables d'environnement (aucun secret reel)
└── README.md                   Documentation d'installation et d'exploitation"""

    para = doc.add_paragraph()
    run = para.add_run(tree)
    run.font.name = "Consolas"
    run.font.size = Pt(8.3)
    para.paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# 21. Technologies et dependances
# ---------------------------------------------------------------------------

def add_section_21(doc):
    doc.add_heading("21. Technologies et dependances", level=1)
    rows = [
        ["Backend", "Spring Boot", "4.1.0 (pom.xml)", "Framework applicatif REST/SSE"],
        ["Backend", "Java", "17", "Langage / JDK"],
        ["Backend", "Spring Data JPA + Hibernate", "heritee du parent Spring Boot", "Persistance ORM"],
        ["Backend", "Flyway (+ flyway-database-postgresql)", "heritee du parent", "Migrations de schema"],
        ["Backend", "Spring WebFlux (WebClient)", "heritee du parent", "Appels reactifs sortants (DLP, LiteLLM, Keycloak Admin)"],
        ["Backend", "spring-boot-starter-oauth2-resource-server", "heritee du parent", "Validation JWT Keycloak"],
        ["Backend", "PostgreSQL JDBC Driver", "heritee du parent", "Pilote base de donnees"],
        ["Frontend", "React / React DOM", "^19.2.7", "Bibliotheque d'interface"],
        ["Frontend", "Vite", "^8.1.1", "Bundler / serveur de developpement"],
        ["Frontend", "keycloak-js", "^26.2.4", "Adaptateur d'authentification Keycloak"],
        ["Frontend", "react-markdown / remark-gfm / rehype-sanitize", "^10.1.0 / ^4.0.1 / ^6.0.0", "Rendu Markdown securise des messages"],
        ["Frontend", "react-syntax-highlighter", "^16.1.1", "Coloration syntaxique du code"],
        ["Frontend", "mammoth", "^1.12.0", "Previsualisation .docx cote client"],
        ["Frontend", "vitest", "^4.0.15", "Tests unitaires"],
        ["Service DLP", "FastAPI", "0.116.1", "Framework web"],
        ["Service DLP", "Uvicorn", "0.35.0", "Serveur ASGI"],
        ["Service DLP", "Presidio Analyzer", "2.2.359", "Detection NER de donnees sensibles"],
        ["Service DLP", "spaCy (+ modeles en/fr)", "3.7.5", "Moteur NLP sous-jacent a Presidio"],
        ["Service DLP", "PyMuPDF / python-docx / python-pptx / openpyxl", "1.26.3 / 1.2.0 / 1.0.2 / 3.1.5", "Extraction de contenu par format de document"],
        ["Service DLP", "pytesseract + Tesseract OCR", "0.3.13 (+ moteur systeme)", "OCR (images et documents scannes)"],
        ["Proxy LLM", "LiteLLM", "image :latest (non figee — voir section 22)", "Normalisation des appels multi-fournisseurs"],
        ["Authentification", "Keycloak", "26.7.0", "Serveur OIDC / gestion des identites"],
        ["Base de donnees", "PostgreSQL", "16-alpine (x2 instances)", "Persistance applicative et Keycloak"],
        ["Infrastructure", "Docker Compose", "format docker-compose.yml", "Orchestration des services"],
        ["CI", "GitHub Actions", ".github/workflows/ci.yml", "Integration continue (4 jobs)"],
    ]
    add_table(doc, ["Composant", "Technologie", "Version", "Role"], rows, col_widths=[2.5, 5.0, 4.0, 4.5])


# ---------------------------------------------------------------------------
# 22. Limitations et etat actuel
# ---------------------------------------------------------------------------

def add_section_22(doc):
    doc.add_heading("22. Limitations et etat actuel", level=1)
    p(doc, "Cette section recense uniquement des elements directement observables dans le "
           "repository (code, configuration, commentaires, README) — aucune supposition.")

    doc.add_heading("22.1 Code non cable / vestiges", level=2)
    bullets(doc, [
        "ContentNotAllowedException et son gestionnaire ChatExceptionHandler sont definis mais jamais leves par aucun code du backend.",
        "AddBannedWordRequest et AddLlmRestrictionRequest (DTO valides) sont declares mais les controllers correspondants utilisent en realite des Map<String,String> brutes sans validation typee — refactorisation probablement inachevee.",
        "DemoUserProvider (service) n'est reference par aucun controller observe — vestige d'un mode demo anterieur, remplace par la resolution JWT reelle (CurrentUserService).",
        "StartupDataFixer insere un utilisateur demo-user et un utilisateur de test a UUID fixe au demarrage, mais uniquement en profil dev (@Profile(\"dev\")).",
        "backend/src/main/resources/db/migration_backup/ contient deux anciens scripts de migration qui ne sont pas dans le chemin Flyway actif — sans effet sur le schema reel.",
        "AdminSidebar et AdminShell (frontend, AdminComponents.jsx) ne sont plus importes nulle part depuis que la navigation admin a ete integree au composant Sidebar.jsx principal — code probablement mort.",
        "frontend/public/assets/admin sidebar/ (avec un espace dans le nom) contient des icones qui ne semblent referencees par aucun composant ; les icones reellement utilisees sont les fichiers plats /assets/admin-*.png — dossier probablement a nettoyer (etat non commite au moment de l'analyse, visible dans git status).",
    ])

    doc.add_heading("22.2 Incoherences documentaires relevaées", level=2)
    bullets(doc, [
        "litellm/README.md documente encore secure-groq -> groq/llama-3.1-8b-instant, alors que litellm/config.yaml (source de verite) utilise groq/openai/gpt-oss-20b depuis la migration V22.",
        "database/README.md ne liste que les 5 tables creees jusqu'a la migration V4 (fournisseur_llm, modele_llm, utilisateur, conversation, message) et omet attachment, audit_logs, filtered_messages et les 5 tables de permissions ajoutees en V15 — documentation partiellement obsolete par rapport au schema reel (V22).",
    ])

    doc.add_heading("22.3 Points de securite/robustesse a surveiller", level=2)
    bullets(doc, [
        "L'endpoint POST /analyse du service DLP (texte JSON pur) ne comporte pas de bloc try/except explicite autour de l'appel au pipeline — une exception interne inattendue y produirait une erreur HTTP 500 generique plutot que le format BLOCK habituel du service (le comportement observable de bout en bout reste neanmoins fail-closed cote backend, qui traite toute erreur de transport comme une indisponibilite).",
        "Le realm Keycloak importe n'expose pas explicitement l'attribut PKCE sur le client synapse-client ; PKCE S256 est actif via le comportement par defaut de keycloak-js plutot que par une politique explicite du realm. Le client conserve egalement directAccessGrantsEnabled=true.",
        "L'image Docker LiteLLM est reference en tag :latest (non figee) dans docker-compose.yml — le README du projet signale lui-meme qu'il faudra figer une version precise apres validation.",
        "Le middleware de controle de taille du service DLP (base sur l'en-tete Content-Length) est documente dans le code comme contournable via un transfert chunked sans Content-Length ; un second filet (lecture bornee du flux) attenue ce risque.",
        "Aucune contrainte CHECK SQL n'encadre attachment.dlp_decision/extraction_status ni filtered_messages.action/request_status — validation uniquement applicative.",
        "Les regles de detection specifiques au Maroc (CIN, RIB, IBAN, telephone) sont explicitement qualifiees de « regles MVP personnalisees » dans le README du service DLP, a revalider avant un usage en production reelle.",
    ])

    doc.add_heading("22.4 Restrictions assumees (choix produit, pas des bugs)", level=2)
    bullets(doc, [
        "Les entites Presidio de type LOCATION sont volontairement toujours ignorees par le detecteur DLP (desactivation deliberee, documentee dans le code et le README du service).",
        "Le NER Presidio (spaCy) ne s'execute que sur le francais et l'anglais ; un texte detecte comme arabe n'est pas soumis a l'analyse Presidio (l'OCR et les regles regex restent actifs sur l'arabe).",
        "L'adaptateur de detection par modele transformer existe dans le code mais est desactive par defaut et sa dependance n'est pas installee par defaut.",
        "Le README du service DLP indique explicitement : « This is not a production banking DLP system. »",
    ])

    doc.add_heading("22.5 Absence de couverture de tests mesuree", level=2)
    p(doc, "Aucun rapport de couverture (JaCoCo, coverage.py, lcov) n'est present dans le depot. Les "
           "trois suites de tests (backend, frontend, DLP) s'executent toutes dans le pipeline "
           "d'integration continue, mais aucun pourcentage de couverture n'est mesure ni "
           "communique.")


# ---------------------------------------------------------------------------
# 23. Annexes
# ---------------------------------------------------------------------------

def add_section_23(doc):
    doc.add_heading("23. Annexes", level=1)

    doc.add_heading("23.1 Matrice des roles Keycloak", level=2)
    add_table(doc, ["Role", "Portee observee dans le code"], [
        ["ADMIN", "Acces a l'ensemble des endpoints /api/admin/** (@PreAuthorize hasRole('ADMIN')) ; seul role pouvant contourner une restriction de role (jamais une restriction personnelle) sur l'acces aux modeles."],
        ["INTERN", "Role realm applicatif (« Internal Synapse user »), utilisable comme cible de restrictions de modeles/mots bannis par role."],
        ["EXTERN", "Role realm applicatif (« External Synapse user »), meme usage qu'INTERN."],
        ["USER", "Role realm de base pour un utilisateur standard."],
    ], col_widths=[2.5, 13.5])

    doc.add_heading("23.2 Liste des conteneurs Docker Compose", level=2)
    add_table(doc, ["Conteneur", "Port hote"], [
        ["postgres", "5433"],
        ["secure-llm-litellm (litellm)", "4000"],
        ["dlp-service", "127.0.0.1:8000"],
        ["keycloak-db", "aucun (interne)"],
        ["keycloak", "8080"],
        ["keycloak-provisioner", "aucun (job one-shot)"],
    ], col_widths=[6.0, 10.0])

    doc.add_heading("23.3 Ports par defaut de l'ensemble du systeme", level=2)
    add_table(doc, ["Service", "Port"], [
        ["Frontend (Vite dev server)", "5173"],
        ["Backend (API sous /api)", "8081"],
        ["Service DLP", "8000"],
        ["LiteLLM", "4000"],
        ["Keycloak", "8080"],
        ["PostgreSQL applicatif (depuis l'hote)", "5433"],
        ["PostgreSQL applicatif (reseau Docker interne)", "5432"],
    ], col_widths=[9.0, 7.0])

    doc.add_heading("23.4 Variables de configuration cles (recapitulatif)", level=2)
    p(doc, "Voir le detail complet section 15. Rappel des variables les plus sensibles (valeurs "
           "toujours masquees) : LITELLM_MASTER_KEY, DLP_ADMIN_KEY, POSTGRES_PASSWORD, "
           "SPRING_DATASOURCE_PASSWORD, KEYCLOAK_DB_PASSWORD, KEYCLOAK_ADMIN_PASSWORD, "
           "KEYCLOAK_DEMO_PASSWORD, GATEWAY_ADMIN_CLIENT_SECRET, OPENAI_API_KEY, GROQ_API_KEY, "
           "GEMINI_API_KEY, MISTRAL_API_KEY, ANTHROPIC_API_KEY — toutes definies uniquement dans le "
           "fichier .env local, non versionne.")

    doc.add_heading("23.5 Matrice synthetique des endpoints par niveau d'acces", level=2)
    add_table(doc, ["Niveau d'acces", "Nombre d'endpoints (approx.)", "Exemples"], [
        ["Public (permitAll)", "1", "/actuator/health"],
        ["JWT valide (sans role specifique)", "~15", "/api/health, /api/chat, /api/conversations/**, /api/attachments/**, /api/models/**"],
        ["JWT + role ADMIN", "~30", "/api/admin/metrics/**, /api/admin/models/**, /api/admin/permissions/**, /api/admin/audit, /api/admin/filtered-messages, /api/admin/keycloak/**"],
    ], col_widths=[5.0, 4.5, 6.5])


