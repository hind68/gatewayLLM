#!/usr/bin/env python3
"""Generateur du Dossier d'Architecture Technique (DAT) - Secure LLM Gateway.

Regenere documents/DAT.docx a partir des donnees codees ci-dessous. Ces
donnees doivent rester alignees avec le code source reel du projet (voir
regle permanente dans CLAUDE.md). Pour mettre a jour le DAT : modifier les
sections/donnees concernees dans ce script, regenerer les diagrammes si
necessaire (fonctions build_*_diagram), puis relancer:

    python documents/generate_dat.py

Le script (re)genere aussi les images de diagrammes dans documents/diagrams/.
"""

import datetime
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parent
DIAGRAMS_DIR = ROOT / "diagrams"
DIAGRAMS_DIR.mkdir(exist_ok=True)
OUTPUT_PATH = ROOT / "DAT.docx"

DOC_VERSION = "1.0"
DOC_DATE = datetime.date(2026, 8, 28)
PROJECT_NAME = "Secure LLM Gateway"

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x2E, 0x6F, 0x9E)
GREY = RGBColor(0x59, 0x59, 0x59)
WARN = RGBColor(0xB2, 0x3A, 0x48)
GREEN = RGBColor(0x2E, 0x7D, 0x4F)
LIGHT_GREY = "F2F2F2"
HEADER_BLUE = "1B2A4A"


# ---------------------------------------------------------------------------
# Low level helpers
# ---------------------------------------------------------------------------

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def add_table(doc, headers, rows, col_widths=None, header_color=HEADER_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_background(hdr_cells[i], header_color)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=9.5)
            if r_idx % 2 == 1:
                set_cell_background(cells[i], LIGHT_GREY)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    return p


def add_figure(doc, image_path, caption, width_cm=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_page_break(doc):
    doc.add_page_break()


def add_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Cliquez ici pour mettre a jour la table des matieres (F9 dans Word)."
    fldChar2.append(fldChar3)
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------

def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.page_break_before = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = ACCENT
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = GREY
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    if "Caption" in [s.name for s in doc.styles]:
        cap = doc.styles["Caption"]
        cap.font.size = Pt(9)
        cap.font.italic = True

    # List Bullet size
    for name in ("List Bullet", "List Number"):
        try:
            st = doc.styles[name]
            st.font.size = Pt(10.5)
            st.font.name = "Calibri"
        except KeyError:
            pass


def setup_sections(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def add_running_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = ""
    run = p.add_run(f"{PROJECT_NAME} — Dossier d'Architecture Technique")
    run.font.size = Pt(8.5)
    run.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_run = fp.add_run(f"DAT — v{DOC_VERSION} — Confidentiel     ")
    left_run.font.size = Pt(8.5)
    left_run.font.color.rgb = GREY
    add_page_number_field(fp)


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def add_cover_page(doc):
    section = doc.sections[0]
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DOSSIER D'ARCHITECTURE TECHNIQUE")
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("(DAT)")
    run2.font.size = Pt(20)
    run2.font.color.rgb = ACCENT

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(PROJECT_NAME)
    run3.font.size = Pt(22)
    run3.font.bold = True
    run3.font.color.rgb = RGBColor(0, 0, 0)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Passerelle securisee d'acces aux modeles de langage (LLM)")
    run4.font.size = Pt(13)
    run4.italic = True
    run4.font.color.rgb = GREY

    for _ in range(8):
        doc.add_paragraph()

    meta = [
        ("Version du document", DOC_VERSION),
        ("Date de derniere mise a jour", DOC_DATE.strftime("%d/%m/%Y")),
        ("Statut", "Vivant — maintenu en continu"),
        ("Classification", "Usage interne / academique"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (k, v) in enumerate(meta):
        row = table.rows[i]
        row.cells[0].width = Cm(6.5)
        row.cells[1].width = Cm(7.5)
        set_cell_text(row.cells[0], k, bold=True, size=11)
        set_cell_text(row.cells[1], v, size=11)
        set_cell_background(row.cells[0], LIGHT_GREY)

    for _ in range(6):
        doc.add_paragraph()

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "Ce document decrit l'architecture reelle du systeme telle qu'observee dans le\n"
        "code source, la configuration et l'infrastructure du projet."
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = GREY

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Main assembly (content added by build script sections)
# ---------------------------------------------------------------------------

def new_document():
    doc = Document()
    setup_sections(doc)
    setup_styles(doc)
    add_cover_page(doc)
    add_running_header_footer(doc)

    doc.add_heading("Table des matieres", level=1)
    doc.paragraphs[-1].paragraph_format.page_break_before = False
    add_toc(doc)
    return doc


if __name__ == "__main__":
    doc = new_document()

    import dat_content as C
    for i in range(1, 24):
        fn = getattr(C, f"add_section_{i:02d}", None)
        if fn:
            fn(doc)

    doc.save(OUTPUT_PATH)
    print(f"DAT genere: {OUTPUT_PATH}")
